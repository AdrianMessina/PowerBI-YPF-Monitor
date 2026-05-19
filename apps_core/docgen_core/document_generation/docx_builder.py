"""
Word Document Builder

Main orchestrator for building complete Word documentation from Power BI metadata.
"""

from docx import Document
from pathlib import Path
from datetime import datetime
from typing import Optional
import logging

from apps_core.docgen_core.core.models.report_metadata import ReportMetadata
from apps_core.docgen_core.core.validators.validation_report import ValidationReport
from apps_core.docgen_core.visualization.er_diagram_generator import ERDiagramGenerator
from .template_handler import TemplateHandler
from .section_generators import (
    CoverPageGenerator,
    ExecutiveSummaryGenerator,
    DataModelSectionGenerator,
    RelationshipsSectionGenerator,
    DAXMeasuresSectionGenerator,
    SecuritySectionGenerator,
    VisualizationsSectionGenerator,
    ValidationSectionGenerator,
    AppendixSectionGenerator
)

logger = logging.getLogger(__name__)


class DocxBuilder:
    """Main builder for Power BI documentation Word documents"""

    def __init__(self, template_path: str):
        """
        Initialize the document builder

        Args:
            template_path: Path to corporate Word template
        """
        self.template_path = Path(template_path)
        self.doc: Optional[Document] = None
        self.template_handler: Optional[TemplateHandler] = None
        self.metadata: Optional[ReportMetadata] = None
        self.validation_report: Optional[ValidationReport] = None
        self.er_diagram_generator: Optional[ERDiagramGenerator] = None

    def build(self, metadata: ReportMetadata, validation_report: ValidationReport,
              er_diagram_generator: Optional[ERDiagramGenerator] = None,
              progress_callback=None) -> str:
        """
        Build complete Word document

        Args:
            metadata: Report metadata
            validation_report: Validation report
            er_diagram_generator: ER diagram generator (optional)
            progress_callback: Optional callback function(step: int, message: str)

        Returns:
            Path to generated document
        """
        logger.info("Starting document generation")

        # Initialize document
        self._progress(progress_callback, 5, "Inicializando documento desde plantilla...")
        self._initialize_document()

        # Store data
        self.metadata = metadata
        self.validation_report = validation_report
        self.er_diagram_generator = er_diagram_generator

        try:
            # Generate all sections in order
            self._progress(progress_callback, 10, "Generando portada...")
            self._generate_cover_page()

            self._progress(progress_callback, 15, "Generando resumen ejecutivo...")
            self._generate_executive_summary()

            self._progress(progress_callback, 20, "Agregando tabla de contenidos...")
            self._generate_table_of_contents()

            self._progress(progress_callback, 30, "Documentando modelo de datos...")
            self._generate_data_model_section()

            self._progress(progress_callback, 45, "Documentando relaciones...")
            self._generate_relationships_section()

            self._progress(progress_callback, 55, "Documentando medidas DAX...")
            self._generate_dax_section()

            self._progress(progress_callback, 70, "Documentando configuración de seguridad...")
            self._generate_security_section()

            self._progress(progress_callback, 80, "Documentando visualizaciones...")
            self._generate_visualizations_section()

            self._progress(progress_callback, 90, "Generando reporte de validación...")
            self._generate_validation_section()

            self._progress(progress_callback, 95, "Agregando apéndice...")
            self._generate_appendix()

            # Save document
            self._progress(progress_callback, 98, "Guardando documento...")
            output_path = self._get_output_path()
            self.doc.save(str(output_path))

            self._progress(progress_callback, 100, "¡Documento generado exitosamente!")
            logger.info(f"Document generated successfully: {output_path}")

            return str(output_path)

        except Exception as e:
            logger.error(f"Error generating document: {e}", exc_info=True)
            raise

    def _initialize_document(self) -> None:
        """Initialize document from template"""
        if not self.template_path.exists():
            logger.warning(f"Template not found at {self.template_path}, using blank document")
            self.doc = Document()
        else:
            self.doc = Document(str(self.template_path))

        self.template_handler = TemplateHandler(self.doc)

    def _generate_cover_page(self) -> None:
        """Generate cover page"""
        generator = CoverPageGenerator(self.doc, self.metadata, self.template_handler)
        generator.generate()

    def _generate_executive_summary(self) -> None:
        """Generate executive summary"""
        generator = ExecutiveSummaryGenerator(
            self.doc, self.metadata, self.validation_report, self.template_handler
        )
        generator.generate()

    def _generate_table_of_contents(self) -> None:
        """Generate table of contents placeholder"""
        self.doc.add_page_break()
        heading = self.doc.add_heading('Tabla de Contenidos', level=1)

        # Add instruction for updating TOC
        paragraph = self.doc.add_paragraph(
            'Haga clic derecho aquí y seleccione "Actualizar campos" para generar la tabla de contenidos.'
        )
        paragraph.italic = True

        # Note: Automatic TOC generation requires Word automation or complex XML
        # Users can easily generate it in Word via Insert > Table of Contents

        self.doc.add_page_break()

    def _generate_data_model_section(self) -> None:
        """Generate data model section"""
        generator = DataModelSectionGenerator(
            self.doc, self.metadata, self.template_handler, self.er_diagram_generator
        )
        generator.generate()

    def _generate_relationships_section(self) -> None:
        """Generate relationships section"""
        generator = RelationshipsSectionGenerator(
            self.doc, self.metadata, self.template_handler
        )
        generator.generate()

    def _generate_dax_section(self) -> None:
        """Generate DAX measures section"""
        generator = DAXMeasuresSectionGenerator(
            self.doc, self.metadata, self.template_handler
        )
        generator.generate()

    def _generate_security_section(self) -> None:
        """Generate security section"""
        generator = SecuritySectionGenerator(
            self.doc, self.metadata, self.template_handler
        )
        generator.generate()

    def _generate_visualizations_section(self) -> None:
        """Generate visualizations section"""
        generator = VisualizationsSectionGenerator(
            self.doc, self.metadata, self.template_handler
        )
        generator.generate()

    def _generate_validation_section(self) -> None:
        """Generate validation section"""
        generator = ValidationSectionGenerator(
            self.doc, self.metadata, self.validation_report, self.template_handler
        )
        generator.generate()

    def _generate_appendix(self) -> None:
        """Generate appendix"""
        generator = AppendixSectionGenerator(
            self.doc, self.metadata, self.template_handler
        )
        generator.generate()

    def _get_output_path(self) -> Path:
        """
        Generate output filename with timestamp

        Returns:
            Path to output file
        """
        # Create output directory if it doesn't exist
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)

        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_name = self.metadata.report_name.replace(" ", "_")
        # Remove invalid filename characters
        report_name = "".join(c for c in report_name if c.isalnum() or c in ('_', '-'))
        filename = f"PowerBI_Documentation_{report_name}_{timestamp}.docx"

        return output_dir / filename

    def _progress(self, callback, step: int, message: str) -> None:
        """
        Call progress callback if provided

        Args:
            callback: Progress callback function
            step: Progress percentage (0-100)
            message: Progress message
        """
        if callback:
            callback(step, message)
