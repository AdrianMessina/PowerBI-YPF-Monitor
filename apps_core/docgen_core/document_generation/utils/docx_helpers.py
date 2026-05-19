"""
Word Document Formatting Utilities

Helper functions for formatting Word documents with python-docx.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Any, Optional


class DocxHelpers:
    """Utilities for formatting Word documents"""

    @staticmethod
    def add_heading(doc: Document, text: str, level: int = 1) -> Any:
        """
        Add formatted heading to document

        Args:
            doc: Document object
            text: Heading text
            level: Heading level (1-9)

        Returns:
            Paragraph object
        """
        heading = doc.add_heading(text, level=level)
        return heading

    @staticmethod
    def add_paragraph(doc: Document, text: str, style: str = 'Normal') -> Any:
        """
        Add formatted paragraph to document

        Args:
            doc: Document object
            text: Paragraph text
            style: Style name (ignored - always creates plain paragraph)

        Returns:
            Paragraph object
        """
        # Don't use styles - create plain paragraph
        paragraph = doc.add_paragraph(text)
        return paragraph

    @staticmethod
    def add_table(doc: Document, data: List[List[str]], headers: List[str],
                  style: str = 'Table Grid') -> Any:
        """
        Add formatted table to document

        Args:
            doc: Document object
            data: List of rows, each row is a list of cell values
            headers: List of column headers
            style: Table style name

        Returns:
            Table object
        """
        # Create table with header row + data rows
        table = doc.add_table(rows=1 + len(data), cols=len(headers))

        # Try to apply style, but don't fail if it doesn't exist
        try:
            table.style = style
        except KeyError:
            # Style doesn't exist, use default or skip
            pass

        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = str(header)
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add data rows
        for row_idx, row_data in enumerate(data, start=1):
            cells = table.rows[row_idx].cells
            for col_idx, value in enumerate(row_data):
                cells[col_idx].text = str(value) if value is not None else ''

        return table

    @staticmethod
    def add_code_block(doc: Document, code: str, language: str = 'dax') -> Any:
        """
        Add code block with monospace font

        Args:
            doc: Document object
            code: Code text
            language: Programming language (for reference, not syntax highlighting)

        Returns:
            Paragraph object
        """
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(code)

        # Apply code formatting
        DocxHelpers.format_code_text(run)

        # Add light background shading
        DocxHelpers._add_shading(paragraph, "F0F0F0")

        return paragraph

    @staticmethod
    def add_colored_badge(paragraph: Any, text: str, color: str) -> Any:
        """
        Add colored badge inline (via highlighting)

        Args:
            paragraph: Paragraph object to add badge to
            text: Badge text
            color: Color name (red, yellow, green, lightGray)

        Returns:
            Run object
        """
        run = paragraph.add_run(f" {text} ")
        run.bold = True
        run.font.size = Pt(9)

        # Map color names to RGB values
        color_map = {
            'red': RGBColor(220, 53, 69),
            'orange': RGBColor(253, 126, 20),
            'yellow': RGBColor(255, 193, 7),
            'green': RGBColor(40, 167, 69),
            'lightGray': RGBColor(108, 117, 125)
        }

        if color.lower() in color_map:
            run.font.color.rgb = color_map[color.lower()]

        return run

    @staticmethod
    def add_page_break(doc: Document) -> None:
        """
        Add page break

        Args:
            doc: Document object
        """
        doc.add_page_break()

    @staticmethod
    def add_bulleted_list(doc: Document, items: List[str]) -> None:
        """
        Add bulleted list

        Args:
            doc: Document object
            items: List of items
        """
        # Create bullets manually without relying on styles
        for item in items:
            paragraph = doc.add_paragraph()
            paragraph.text = f"• {item}"
            # Add indentation
            from docx.shared import Pt
            paragraph.paragraph_format.left_indent = Pt(18)

    @staticmethod
    def format_code_text(run: Any) -> None:
        """
        Apply code formatting to a run

        Args:
            run: Run object
        """
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(51, 51, 51)

    @staticmethod
    def add_emphasis_box(doc: Document, text: str, title: Optional[str] = None) -> Any:
        """
        Add emphasized box for recommendations or warnings

        Args:
            doc: Document object
            text: Content text
            title: Optional title for the box

        Returns:
            Paragraph object
        """
        if title:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(10)

        paragraph = doc.add_paragraph(text)
        # Add border
        DocxHelpers._add_border(paragraph)
        # Add shading
        DocxHelpers._add_shading(paragraph, "E7F3FF")

        return paragraph

    @staticmethod
    def _add_shading(paragraph: Any, color_hex: str) -> None:
        """
        Add shading to paragraph background

        Args:
            paragraph: Paragraph object
            color_hex: Hex color without #
        """
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        paragraph._p.get_or_add_pPr().append(shading_elm)

    @staticmethod
    def _add_border(paragraph: Any, color: str = "4A90E2", width: int = 4) -> None:
        """
        Add border to paragraph

        Args:
            paragraph: Paragraph object
            color: Hex color without #
            width: Border width in points
        """
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(width))
            border.set(qn('w:space'), '1')
            border.set(qn('w:color'), color)
            pBdr.append(border)

        pPr.append(pBdr)

    @staticmethod
    def add_numbered_list(doc: Document, items: List[str]) -> None:
        """
        Add numbered list

        Args:
            doc: Document object
            items: List of items
        """
        # Create numbered list manually without relying on styles
        for idx, item in enumerate(items, 1):
            paragraph = doc.add_paragraph()
            paragraph.text = f"{idx}. {item}"
            # Add indentation
            from docx.shared import Pt
            paragraph.paragraph_format.left_indent = Pt(18)

    @staticmethod
    def add_image(doc: Document, image_path: str, width_inches: float = 6.5,
                  caption: Optional[str] = None) -> Any:
        """
        Add image to document

        Args:
            doc: Document object
            image_path: Path to image file
            width_inches: Image width in inches
            caption: Optional caption text

        Returns:
            Paragraph containing image
        """
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(width_inches))

        if caption:
            # Create caption manually without relying on Caption style
            caption_para = doc.add_paragraph(caption)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption_para.runs:
                run.italic = True
                run.font.size = Pt(9)

        return paragraph

    @staticmethod
    def add_statistics_table(doc: Document, stats: dict, title: Optional[str] = None) -> Any:
        """
        Add a 2-column statistics table (Label | Value)

        Args:
            doc: Document object
            stats: Dictionary of label: value pairs
            title: Optional title

        Returns:
            Table object
        """
        if title:
            doc.add_heading(title, level=3)

        table = doc.add_table(rows=len(stats), cols=2)
        table.style = 'Table Grid'  # Use standard style

        for idx, (label, value) in enumerate(stats.items()):
            cells = table.rows[idx].cells
            cells[0].text = str(label)
            cells[1].text = str(value)

            # Make label bold
            for paragraph in cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        return table
