"""
Validation Section Generator

Genera reporte de validación con puntuación de calidad, problemas y recomendaciones.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from typing import Any
import logging

from ..utils.docx_helpers import DocxHelpers

logger = logging.getLogger(__name__)


class ValidationSectionGenerator:
    """Genera sección de reporte de validación"""

    def __init__(self, doc: Document, metadata: Any, validation_report: Any,
                 template_handler: Any):
        """
        Initialize validation section generator

        Args:
            doc: Document object
            metadata: ReportMetadata object
            validation_report: ValidationReport object
            template_handler: TemplateHandler object
        """
        self.doc = doc
        self.metadata = metadata
        self.validation_report = validation_report
        self.template_handler = template_handler

    def generate(self) -> None:
        """Generate validation section"""
        logger.info("Generating validation section")

        DocxHelpers.add_heading(self.doc, 'Reporte de Validación', level=1)

        if not self.validation_report:
            DocxHelpers.add_paragraph(self.doc, "No hay reporte de validación disponible.")
            return

        # Quality score
        self._add_quality_score()

        # Issue summary
        self._add_issue_summary()

        # Issues by severity
        self._add_issues_by_severity()

        # Recommendations
        self._add_recommendations()

        logger.info("Validation section generated successfully")

    def _add_quality_score(self) -> None:
        """Add quality score with visual representation"""
        if not hasattr(self.validation_report, 'quality_score'):
            return

        score = self.validation_report.quality_score

        DocxHelpers.add_heading(self.doc, 'Puntuación de Calidad', level=2)

        # Score paragraph with color
        score_para = self.doc.add_paragraph()
        score_run = score_para.add_run(f"{score}/100")
        score_run.bold = True
        score_run.font.size = Pt(24)

        # Color based on score
        if score >= 90:
            score_run.font.color.rgb = RGBColor(16, 185, 129)  # Green
        elif score >= 75:
            score_run.font.color.rgb = RGBColor(245, 158, 11)  # Yellow
        elif score >= 60:
            score_run.font.color.rgb = RGBColor(253, 126, 20)  # Orange
        else:
            score_run.font.color.rgb = RGBColor(220, 53, 69)  # Red

        # Interpretation
        if score >= 90:
            interpretation = "Excelente - El reporte sigue las mejores prácticas con problemas mínimos."
        elif score >= 75:
            interpretation = "Bueno - El reporte está bien estructurado con algunas mejoras menores necesarias."
        elif score >= 60:
            interpretation = "Regular - Varios problemas requieren atención para mejorar la calidad del reporte."
        else:
            interpretation = "Necesita Mejoras - Múltiples problemas críticos deben ser atendidos."

        DocxHelpers.add_paragraph(self.doc, interpretation)

    def _add_issue_summary(self) -> None:
        """Add issue counts by severity"""
        if not hasattr(self.validation_report, 'issues'):
            return

        issues = self.validation_report.issues

        DocxHelpers.add_heading(self.doc, 'Resumen de Problemas', level=2)

        # Count by severity
        critical = sum(1 for i in issues if hasattr(i, 'severity') and i.severity == 'CRITICAL')
        errors = sum(1 for i in issues if hasattr(i, 'severity') and i.severity == 'ERROR')
        warnings = sum(1 for i in issues if hasattr(i, 'severity') and i.severity == 'WARNING')
        info = sum(1 for i in issues if hasattr(i, 'severity') and i.severity == 'INFO')

        # Create summary table
        data = [
            ['Crítico', str(critical), 'Problemas que requieren atención inmediata'],
            ['Errores', str(errors), 'Problemas que deberían ser corregidos'],
            ['Advertencias', str(warnings), 'Problemas que pueden impactar el rendimiento o usabilidad'],
            ['Información', str(info), 'Recomendaciones de mejores prácticas']
        ]
        headers = ['Severidad', 'Cantidad', 'Descripción']
        DocxHelpers.add_table(self.doc, data, headers, style='Table Grid')

    def _add_issues_by_severity(self) -> None:
        """Add detailed issues grouped by severity"""
        if not hasattr(self.validation_report, 'issues'):
            return

        issues = self.validation_report.issues
        if not issues:
            DocxHelpers.add_heading(self.doc, 'Problemas', level=2)
            DocxHelpers.add_paragraph(self.doc, "¡No se encontraron problemas. El reporte está en excelente condición!")
            return

        DocxHelpers.add_heading(self.doc, 'Problemas Detallados', level=2)

        # Group by severity
        severities = ['CRITICAL', 'ERROR', 'WARNING', 'INFO']
        severity_labels = {
            'CRITICAL': 'CRÍTICOS',
            'ERROR': 'ERRORES',
            'WARNING': 'ADVERTENCIAS',
            'INFO': 'INFORMACIÓN'
        }

        for severity in severities:
            severity_issues = [i for i in issues if hasattr(i, 'severity') and i.severity == severity]

            if not severity_issues:
                continue

            # Severity heading
            DocxHelpers.add_heading(self.doc, f"Problemas {severity_labels[severity]}", level=3)

            # Badge
            badge_para = self.doc.add_paragraph()
            if severity == 'CRITICAL':
                DocxHelpers.add_colored_badge(badge_para, f"{len(severity_issues)} problemas", 'red')
            elif severity == 'ERROR':
                DocxHelpers.add_colored_badge(badge_para, f"{len(severity_issues)} problemas", 'orange')
            elif severity == 'WARNING':
                DocxHelpers.add_colored_badge(badge_para, f"{len(severity_issues)} problemas", 'yellow')
            else:
                DocxHelpers.add_colored_badge(badge_para, f"{len(severity_issues)} elementos", 'lightGray')

            # Document each issue
            for issue in severity_issues:
                self._document_issue(issue, severity)

    def _document_issue(self, issue: Any, severity: str) -> None:
        """
        Document a single issue

        Args:
            issue: Issue object
            severity: Severity level
        """
        # Category and message
        category = getattr(issue, 'category', 'Desconocido')
        message = getattr(issue, 'message', 'No hay mensaje disponible')

        DocxHelpers.add_heading(self.doc, f"[{category}] {message}", level=4)

        # Details
        if hasattr(issue, 'details') and issue.details:
            DocxHelpers.add_paragraph(self.doc, issue.details)

        # Affected objects
        if hasattr(issue, 'affected_objects') and issue.affected_objects:
            DocxHelpers.add_paragraph(self.doc, "Objetos Afectados:")
            DocxHelpers.add_bulleted_list(self.doc, issue.affected_objects)

        # Recommendation
        if hasattr(issue, 'recommendation') and issue.recommendation:
            DocxHelpers.add_emphasis_box(
                self.doc,
                issue.recommendation,
                title="💡 Recomendación"
            )

        # Add spacing
        self.doc.add_paragraph()

    def _add_recommendations(self) -> None:
        """Add recommendations summary"""
        if not hasattr(self.validation_report, 'recommendations'):
            return

        recommendations = self.validation_report.recommendations
        if not recommendations:
            return

        DocxHelpers.add_heading(self.doc, 'Resumen de Recomendaciones', level=2)

        # Group recommendations by priority
        DocxHelpers.add_bulleted_list(self.doc, recommendations)

        # Add general best practices
        DocxHelpers.add_heading(self.doc, 'Mejores Prácticas', level=3)

        best_practices = [
            "Usar diseño de esquema estrella para rendimiento óptimo",
            "Evitar relaciones bidireccionales cuando sea posible",
            "Mantener las medidas DAX simples y enfocadas",
            "Usar variables en DAX para mejorar legibilidad y rendimiento",
            "Documentar medidas complejas con comentarios",
            "Probar seguridad a nivel de fila exhaustivamente antes de despliegue",
            "Revisar regularmente el tamaño del modelo y métricas de rendimiento"
        ]

        DocxHelpers.add_bulleted_list(self.doc, best_practices)
