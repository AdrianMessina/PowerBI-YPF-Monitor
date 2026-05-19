"""
Relationships Section Generator

Genera documentación de relaciones con advertencias para casos especiales.
"""

from docx import Document
from typing import Any
import logging

from ..utils.docx_helpers import DocxHelpers

logger = logging.getLogger(__name__)


class RelationshipsSectionGenerator:
    """Genera sección de documentación de relaciones"""

    def __init__(self, doc: Document, metadata: Any, template_handler: Any):
        """
        Initialize relationships section generator

        Args:
            doc: Document object
            metadata: ReportMetadata object
            template_handler: TemplateHandler object
        """
        self.doc = doc
        self.metadata = metadata
        self.template_handler = template_handler

    def generate(self) -> None:
        """Generate relationships section"""
        logger.info("Generating relationships section")

        DocxHelpers.add_heading(self.doc, 'Relaciones', level=1)

        if not hasattr(self.metadata, 'data_model') or not self.metadata.data_model:
            DocxHelpers.add_paragraph(self.doc, "No hay información del modelo de datos disponible.")
            return

        model = self.metadata.data_model
        relationships = model.relationships if hasattr(model, 'relationships') else []

        if not relationships:
            DocxHelpers.add_paragraph(self.doc, "No hay relaciones definidas en el modelo de datos.")
            return

        # Overview
        self._add_overview(relationships)

        # Relationships table
        self._add_relationships_table(relationships)

        # Special relationships warnings
        self._add_special_relationships(relationships)

        logger.info("Relationships section generated successfully")

    def _add_overview(self, relationships: list) -> None:
        """Add relationships overview"""
        overview_text = (
            f"El modelo de datos contiene {len(relationships)} relaciones que conectan las tablas. "
            "Estas relaciones definen cómo fluyen los datos entre tablas y permiten el filtrado cruzado."
        )
        DocxHelpers.add_paragraph(self.doc, overview_text)

    def _add_relationships_table(self, relationships: list) -> None:
        """Add relationships table"""
        DocxHelpers.add_heading(self.doc, 'Detalles de Relaciones', level=2)

        data = []
        for rel in relationships:
            from_table = getattr(rel, 'from_table', 'Desconocido')
            from_column = getattr(rel, 'from_column', 'Desconocido')
            to_table = getattr(rel, 'to_table', 'Desconocido')
            to_column = getattr(rel, 'to_column', 'Desconocido')

            # Cardinality
            from_card = getattr(rel, 'from_cardinality', 'Desconocido')
            to_card = getattr(rel, 'to_cardinality', 'Desconocido')
            cardinality = f"{from_card}:{to_card}"

            # Direction
            direction = getattr(rel, 'cross_filter_direction', 'Single')
            direction_es = 'Única' if direction == 'Single' else 'Ambas' if direction == 'Both' else direction

            # Active
            active = 'Sí' if getattr(rel, 'is_active', True) else 'No'

            row = [from_table, from_column, '→', to_table, to_column, cardinality, direction_es, active]
            data.append(row)

        headers = ['Tabla Origen', 'Columna Origen', '', 'Tabla Destino', 'Columna Destino', 'Cardinalidad', 'Dirección', 'Activa']
        DocxHelpers.add_table(self.doc, data, headers, style='Table Grid')

    def _add_special_relationships(self, relationships: list) -> None:
        """Add warnings for special relationships"""
        # Check for bidirectional
        bidirectional = [
            r for r in relationships
            if hasattr(r, 'cross_filter_direction') and r.cross_filter_direction == 'Both'
        ]

        # Check for many-to-many
        many_to_many = [
            r for r in relationships
            if hasattr(r, 'from_cardinality') and hasattr(r, 'to_cardinality')
            and r.from_cardinality == 'Many' and r.to_cardinality == 'Many'
        ]

        # Check for inactive
        inactive = [
            r for r in relationships
            if hasattr(r, 'is_active') and not r.is_active
        ]

        if not (bidirectional or many_to_many or inactive):
            return

        DocxHelpers.add_heading(self.doc, 'Relaciones Especiales', level=2)

        if bidirectional:
            DocxHelpers.add_heading(self.doc, 'Relaciones Bidireccionales', level=3)
            DocxHelpers.add_emphasis_box(
                self.doc,
                f"Se encontraron {len(bidirectional)} relaciones bidireccionales. Estas pueden impactar el rendimiento y causar rutas de consulta ambiguas. Revisar cuidadosamente.",
                title="⚠️ Advertencia"
            )
            items = [f"{getattr(r, 'from_table', '?')} ↔ {getattr(r, 'to_table', '?')}" for r in bidirectional]
            DocxHelpers.add_bulleted_list(self.doc, items)

        if many_to_many:
            DocxHelpers.add_heading(self.doc, 'Relaciones Muchos a Muchos', level=3)
            DocxHelpers.add_emphasis_box(
                self.doc,
                f"Se encontraron {len(many_to_many)} relaciones muchos a muchos. Estas requieren una tabla puente para rendimiento óptimo.",
                title="ℹ️ Información"
            )
            items = [f"{getattr(r, 'from_table', '?')} ↔ {getattr(r, 'to_table', '?')}" for r in many_to_many]
            DocxHelpers.add_bulleted_list(self.doc, items)

        if inactive:
            DocxHelpers.add_heading(self.doc, 'Relaciones Inactivas', level=3)
            DocxHelpers.add_paragraph(
                self.doc,
                f"Se encontraron {len(inactive)} relaciones inactivas. Estas típicamente se usan en funciones DAX USERELATIONSHIP()."
            )
            items = [f"{getattr(r, 'from_table', '?')} → {getattr(r, 'to_table', '?')}" for r in inactive]
            DocxHelpers.add_bulleted_list(self.doc, items)
