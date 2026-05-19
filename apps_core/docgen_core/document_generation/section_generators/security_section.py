"""
Security Section Generator

Genera documentación de configuración de seguridad (RLS, OLS).
"""

from docx import Document
from typing import Any
import logging

from ..utils.docx_helpers import DocxHelpers

logger = logging.getLogger(__name__)


class SecuritySectionGenerator:
    """Genera sección de configuración de seguridad"""

    def __init__(self, doc: Document, metadata: Any, template_handler: Any):
        """
        Initialize security section generator

        Args:
            doc: Document object
            metadata: ReportMetadata object
            template_handler: TemplateHandler object
        """
        self.doc = doc
        self.metadata = metadata
        self.template_handler = template_handler

    def generate(self) -> None:
        """Generate security section"""
        logger.info("Generating security section")

        DocxHelpers.add_heading(self.doc, 'Configuración de Seguridad', level=1)

        if not hasattr(self.metadata, 'security') or not self.metadata.security:
            DocxHelpers.add_paragraph(self.doc, "No se encontró configuración de seguridad en este reporte.")
            return

        security = self.metadata.security

        # RLS Configuration
        if hasattr(security, 'roles') and security.roles:
            self._add_rls_configuration(security.roles)
        else:
            DocxHelpers.add_paragraph(self.doc, "No hay roles de seguridad a nivel de fila (RLS) definidos.")

        # OLS Configuration
        if hasattr(security, 'ols_rules') and security.ols_rules:
            self._add_ols_configuration(security.ols_rules)

        # Security best practices
        self._add_security_best_practices()

        logger.info("Security section generated successfully")

    def _add_rls_configuration(self, roles: list) -> None:
        """Add RLS configuration documentation"""
        DocxHelpers.add_heading(self.doc, 'Seguridad a Nivel de Fila (RLS)', level=2)

        overview_text = (
            f"El reporte implementa seguridad a nivel de fila con {len(roles)} rol(es). "
            "RLS restringe el acceso a datos para usuarios basándose en su membresía de rol."
        )
        DocxHelpers.add_paragraph(self.doc, overview_text)

        # Document each role
        for role in roles:
            self._document_role(role)

    def _document_role(self, role: Any) -> None:
        """
        Document a single RLS role

        Args:
            role: Role object
        """
        # Role name
        role_name = getattr(role, 'name', 'Desconocido')
        DocxHelpers.add_heading(self.doc, role_name, level=3)

        # Description
        if hasattr(role, 'description') and role.description:
            DocxHelpers.add_paragraph(self.doc, role.description)

        # Members (if available)
        if hasattr(role, 'members') and role.members:
            DocxHelpers.add_paragraph(self.doc, "Miembros:")
            DocxHelpers.add_bulleted_list(self.doc, role.members)

        # Table permissions
        if hasattr(role, 'table_permissions') and role.table_permissions:
            DocxHelpers.add_heading(self.doc, 'Permisos de Tabla', level=4)

            data = []
            for perm in role.table_permissions:
                table = getattr(perm, 'table', 'Desconocido')
                filter_expr = getattr(perm, 'filter_expression', '')
                description = getattr(perm, 'description', '')

                row = [table, filter_expr[:100] if filter_expr else '-', description[:50] if description else '-']
                data.append(row)

            headers = ['Tabla', 'Expresión de Filtro', 'Descripción']
            DocxHelpers.add_table(self.doc, data, headers, style='Table Grid')

            # Show full expressions for each
            for perm in role.table_permissions:
                table = getattr(perm, 'table', 'Desconocido')
                filter_expr = getattr(perm, 'filter_expression', '')

                if filter_expr and len(filter_expr) > 100:
                    DocxHelpers.add_heading(self.doc, f"Filtro Completo para {table}", level=5)
                    DocxHelpers.add_code_block(self.doc, filter_expr, language='dax')

        # Add spacing
        self.doc.add_paragraph()

    def _add_ols_configuration(self, ols_rules: list) -> None:
        """Add OLS configuration documentation"""
        DocxHelpers.add_heading(self.doc, 'Seguridad a Nivel de Objeto (OLS)', level=2)

        overview_text = (
            f"El reporte implementa seguridad a nivel de objeto con {len(ols_rules)} regla(s). "
            "OLS controla la visibilidad de tablas, columnas o medidas específicas."
        )
        DocxHelpers.add_paragraph(self.doc, overview_text)

        # Document OLS rules
        data = []
        for rule in ols_rules:
            role = getattr(rule, 'role', 'Desconocido')
            object_type = getattr(rule, 'object_type', 'Desconocido')
            object_name = getattr(rule, 'object_name', 'Desconocido')
            permission = getattr(rule, 'permission', 'Desconocido')

            row = [role, object_type, object_name, permission]
            data.append(row)

        headers = ['Rol', 'Tipo de Objeto', 'Nombre de Objeto', 'Permiso']
        DocxHelpers.add_table(self.doc, data, headers, style='Table Grid')

    def _add_security_best_practices(self) -> None:
        """Add security best practices"""
        DocxHelpers.add_heading(self.doc, 'Mejores Prácticas de Seguridad', level=2)

        best_practices = [
            "Probar los roles RLS exhaustivamente con diferentes contextos de usuario",
            "Usar RLS dinámico basado en identidad de usuario (USERNAME() o USERPRINCIPALNAME())",
            "Evitar usar filtros bidireccionales en tablas con RLS",
            "Documentar las reglas de seguridad claramente para mantenimiento",
            "Auditar regularmente la membresía de roles y permisos",
            "Usar OLS para ocultar columnas sensibles de usuarios no autorizados",
            "Probar el impacto en rendimiento de expresiones de filtro complejas",
            "Considerar usar tablas de seguridad (tablas puente) para escenarios complejos"
        ]

        DocxHelpers.add_bulleted_list(self.doc, best_practices)
