"""
Cover Page Generator

Generates the cover page for Power BI documentation.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Any
import logging

logger = logging.getLogger(__name__)


class CoverPageGenerator:
    """Generates document cover page"""

    def __init__(self, doc: Document, metadata: Any, template_handler: Any):
        """
        Initialize cover page generator

        Args:
            doc: Document object
            metadata: ReportMetadata object
            template_handler: TemplateHandler object
        """
        self.doc = doc
        self.metadata = metadata
        self.template_handler = template_handler

    def generate(self) -> None:
        """Generate cover page"""
        logger.info("Generating cover page")

        # Add title
        self._add_title()

        # Add subtitle
        self._add_subtitle()

        # Add spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Add metadata section
        self._add_metadata_section()

        # Add spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Add generation info
        self._add_generation_info()

        # Page break after cover
        self.doc.add_page_break()

        logger.info("Cover page generated successfully")

    def _add_title(self) -> None:
        """Add document title"""
        title = f"Documentación Power BI"
        title_paragraph = self.doc.add_heading(title, level=0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Apply formatting
        for run in title_paragraph.runs:
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = RGBColor(26, 54, 93)  # YPF Dark Blue

    def _add_subtitle(self) -> None:
        """Add report name as subtitle"""
        subtitle = self.metadata.report_name
        subtitle_paragraph = self.doc.add_heading(subtitle, level=1)
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Apply formatting
        for run in subtitle_paragraph.runs:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0, 168, 89)  # YPF Green

    def _add_metadata_section(self) -> None:
        """Add metadata information section"""
        # Create metadata table
        from ..utils.docx_helpers import DocxHelpers

        metadata_dict = {}

        # Add author if available
        if hasattr(self.metadata, 'author') and self.metadata.author:
            metadata_dict['Autor'] = self.metadata.author

        # Add version if available
        if hasattr(self.metadata, 'version') and self.metadata.version:
            metadata_dict['Versión'] = self.metadata.version

        # Add date
        if hasattr(self.metadata, 'created_date') and self.metadata.created_date:
            metadata_dict['Fecha de Creación'] = self.metadata.created_date
        else:
            metadata_dict['Fecha de Creación'] = datetime.now().strftime('%Y-%m-%d')

        # Add description if available
        if hasattr(self.metadata, 'description') and self.metadata.description:
            metadata_dict['Descripción'] = self.metadata.description

        # Add report ID if available
        if hasattr(self.metadata, 'report_id') and self.metadata.report_id:
            metadata_dict['ID del Reporte'] = self.metadata.report_id

        # Create table
        if metadata_dict:
            # Add section heading
            heading = self.doc.add_heading('Información del Reporte', level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add metadata table
            table = self.doc.add_table(rows=len(metadata_dict), cols=2)
            # Try to apply style, but don't fail if it doesn't exist
            try:
                table.style = 'Table Grid'
            except KeyError:
                pass  # Style doesn't exist, use default

            for idx, (label, value) in enumerate(metadata_dict.items()):
                cells = table.rows[idx].cells
                cells[0].text = label
                cells[1].text = str(value)

                # Make label bold
                for paragraph in cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

                # Wrap long text
                if len(str(value)) > 60:
                    for paragraph in cells[1].paragraphs:
                        paragraph.paragraph_format.line_spacing = 1.15

    def _add_generation_info(self) -> None:
        """Add generation timestamp and tool info"""
        # Add timestamp
        timestamp_para = self.doc.add_paragraph()
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        timestamp_text = f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
        timestamp_run = timestamp_para.add_run(timestamp_text)
        timestamp_run.italic = True
        timestamp_run.font.size = Pt(10)
        timestamp_run.font.color.rgb = RGBColor(102, 102, 102)

        # Add generator info
        generator_para = self.doc.add_paragraph()
        generator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        generator_text = "Generador de Documentación Power BI v3.0"
        generator_run = generator_para.add_run(generator_text)
        generator_run.italic = True
        generator_run.font.size = Pt(10)
        generator_run.font.color.rgb = RGBColor(102, 102, 102)

        # Add company info
        company_para = self.doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        company_text = "YPF S.A. - Equipo de TI y Analítica"
        company_run = company_para.add_run(company_text)
        company_run.bold = True
        company_run.font.size = Pt(11)
        company_run.font.color.rgb = RGBColor(26, 54, 93)
