"""
ER Diagram Embedder

Embeds Entity-Relationship diagrams into Word documents.
"""

from docx import Document
from docx.shared import Inches
from pathlib import Path
import tempfile
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class ERDiagramEmbedder:
    """Handles embedding ER diagrams in Word documents"""

    def __init__(self, doc: Document):
        """
        Initialize ER diagram embedder

        Args:
            doc: Document object
        """
        self.doc = doc

    def embed_diagram(self, er_diagram_generator, width_inches: float = 6.5,
                      caption: str = "Figure 1: Entity-Relationship Diagram") -> bool:
        """
        Embed ER diagram into document

        Args:
            er_diagram_generator: ERDiagramGenerator object
            width_inches: Image width in inches
            caption: Image caption

        Returns:
            True if successful, False otherwise
        """
        try:
            logger.info("Embedding ER diagram")

            # Generate PNG from ER diagram
            temp_path = self._generate_temp_png(er_diagram_generator)

            if not temp_path or not temp_path.exists():
                logger.warning("Could not generate ER diagram PNG")
                return False

            # Add image to document
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(str(temp_path), width=Inches(width_inches))

            # Add caption
            caption_para = self.doc.add_paragraph(caption)
            caption_para.style = 'Caption'

            # Cleanup temp file
            try:
                temp_path.unlink()
            except Exception as e:
                logger.warning(f"Could not delete temp file {temp_path}: {e}")

            logger.info("ER diagram embedded successfully")
            return True

        except Exception as e:
            logger.error(f"Error embedding ER diagram: {e}", exc_info=True)
            return False

    def _generate_temp_png(self, er_diagram_generator) -> Optional[Path]:
        """
        Generate temporary PNG file from ER diagram

        Args:
            er_diagram_generator: ERDiagramGenerator object

        Returns:
            Path to temporary PNG file
        """
        try:
            # Create temp file
            temp_dir = Path(tempfile.gettempdir())
            temp_file = temp_dir / f"er_diagram_{id(er_diagram_generator)}.png"

            # Check if generator has the diagram
            if not hasattr(er_diagram_generator, 'fig'):
                logger.warning("ER diagram generator has no figure")
                return None

            # Export to PNG
            # Try using plotly's write_image first
            try:
                er_diagram_generator.fig.write_image(str(temp_file), width=1200, height=800)
                logger.info(f"ER diagram exported to {temp_file}")
                return temp_file
            except Exception as e:
                logger.warning(f"Could not use plotly write_image: {e}")

                # Fallback: Try using kaleido directly
                try:
                    import plotly.io as pio
                    pio.write_image(er_diagram_generator.fig, str(temp_file),
                                    width=1200, height=800, format='png')
                    logger.info(f"ER diagram exported using kaleido to {temp_file}")
                    return temp_file
                except Exception as e2:
                    logger.warning(f"Could not use kaleido: {e2}")

                    # Last fallback: Try using export_png method if available
                    if hasattr(er_diagram_generator, 'export_png'):
                        try:
                            result_path = er_diagram_generator.export_png(str(temp_file))
                            if result_path and Path(result_path).exists():
                                logger.info(f"ER diagram exported using export_png to {result_path}")
                                return Path(result_path)
                        except Exception as e3:
                            logger.warning(f"Could not use export_png: {e3}")

                    logger.error("All ER diagram export methods failed")
                    return None

        except Exception as e:
            logger.error(f"Error generating temp PNG: {e}", exc_info=True)
            return None

    def embed_diagram_from_file(self, image_path: str, width_inches: float = 6.5,
                                caption: str = "Figure 1: Entity-Relationship Diagram") -> bool:
        """
        Embed ER diagram from existing file

        Args:
            image_path: Path to image file
            width_inches: Image width in inches
            caption: Image caption

        Returns:
            True if successful, False otherwise
        """
        try:
            path = Path(image_path)
            if not path.exists():
                logger.warning(f"Image file not found: {image_path}")
                return False

            # Add image to document
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(str(path), width=Inches(width_inches))

            # Add caption
            caption_para = self.doc.add_paragraph(caption)
            caption_para.style = 'Caption'

            logger.info("ER diagram embedded from file successfully")
            return True

        except Exception as e:
            logger.error(f"Error embedding ER diagram from file: {e}", exc_info=True)
            return False
