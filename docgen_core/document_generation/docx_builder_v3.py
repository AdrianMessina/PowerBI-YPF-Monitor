"""
Word Document Builder V3 - FIXED VERSION
Fills the official corporate template with real PBIP data
CORREGIDO: Usa un método más robusto para insertar contenido
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any, List
import logging

from .intelligent_mapper import IntelligentTemplateMapper

logger = logging.getLogger(__name__)


class DocxBuilderV3:
    """
    Fills official corporate Word template with real Power BI data.
    V3: Fixed insertion method - more robust approach
    """

    def __init__(self, template_path: str):
        """
        Initialize builder with template

        Args:
            template_path: Path to official corporate template
        """
        self.template_path = Path(template_path)
        self.doc: Optional[Document] = None
        self.logger = logging.getLogger(__name__)

    def build(self, metadata: Any, user_inputs: Dict[str, Any],
              er_diagram_path: Optional[str] = None,
              er_image_path: Optional[str] = None,
              visualization_images: Optional[List[str]] = None,
              progress_callback=None,
              sections: Optional[List[str]] = None) -> str:
        """
        Build Word document by filling the template

        Args:
            metadata: Parsed PBIP/TMDL metadata
            user_inputs: User-provided fields (objetivo, alcance, etc.)
            er_diagram_path: Path to auto-generated ER diagram image
            er_image_path: Path to user-uploaded ER model image (takes priority)
            visualization_images: List of paths to user-uploaded visualization images
            progress_callback: Optional callback(step, message)

        Returns:
            Path to generated document
        """
        self.logger.info("Starting document generation V3 with improved filling")

        # All sections by default
        if sections is None:
            sections = [
                "portada", "objetivo", "usuarios", "definiciones",
                "origenes", "filtros", "modelo_er", "visualizaciones",
                "rls", "anexo",
            ]

        def _include(name: str) -> bool:
            return name in sections

        # Load template
        self._progress(progress_callback, 5, "Cargando template corporativo...")
        self._load_template()

        # Generate context
        self._progress(progress_callback, 10, "Generando contexto desde metadata...")
        mapper = IntelligentTemplateMapper(metadata, user_inputs)
        context = mapper.generate_context()

        try:
            # Fill each section (only if selected)
            if _include("portada"):
                self._progress(progress_callback, 15, "Rellenando portada...")
                self._fill_portada(context)
                self._progress(progress_callback, 25, "Rellenando versión del documento...")
                self._fill_version_table(context)

            if _include("objetivo"):
                self._progress(progress_callback, 30, "Rellenando objetivo...")
                self._fill_objective(context)
                self._progress(progress_callback, 35, "Rellenando alcance...")
                self._fill_scope(context)

            if _include("usuarios"):
                self._progress(progress_callback, 40, "Rellenando usuarios...")
                self._fill_users(context)

            if _include("definiciones"):
                self._progress(progress_callback, 50, "Rellenando definiciones (KPIs/Medidas)...")
                self._fill_definitions(context)

            if _include("origenes"):
                self._progress(progress_callback, 60, "Rellenando orígenes de datos...")
                self._fill_data_sources(context)

            if _include("filtros"):
                self._progress(progress_callback, 70, "Rellenando filtros...")
                self._fill_filters(context)

            if _include("modelo_er"):
                self._progress(progress_callback, 80, "Rellenando modelo ER...")
                self._fill_er_model(context, er_diagram_path, er_image_path)

            if _include("visualizaciones"):
                self._progress(progress_callback, 85, "Agregando visualizaciones del reporte...")
                self._fill_visualizations_section(context, visualization_images or [])

            if _include("rls"):
                self._progress(progress_callback, 90, "Rellenando RLS y seguridad...")
                self._fill_rls(context)

            if _include("anexo"):
                self._progress(progress_callback, 95, "Rellenando anexo...")
                self._fill_appendix(context)

            # Save document
            self._progress(progress_callback, 98, "Guardando documento...")
            output_path = self._get_output_path(context)
            self.doc.save(str(output_path))

            self._progress(progress_callback, 100, "¡Documento generado exitosamente!")
            self.logger.info(f"Document generated: {output_path}")

            return str(output_path)

        except Exception as e:
            self.logger.error(f"Error generating document: {e}", exc_info=True)
            raise

    def _load_template(self):
        """Load the corporate template"""
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {self.template_path}")

        self.doc = Document(str(self.template_path))
        self.logger.info(f"Template loaded: {self.template_path}")

    # === SECTION FILLERS ===

    def _replace_text_in_paragraph(self, paragraph, old_text, new_text):
        """
        Replace text in paragraph handling fragmented runs

        Args:
            paragraph: Document paragraph
            old_text: Text to find (e.g., '[Nombre del tablero]')
            new_text: Replacement text (e.g., 'Análisis de ventas')

        Returns:
            True if replacement was made, False otherwise
        """
        # Check if placeholder exists in paragraph
        if old_text not in paragraph.text:
            return False

        # Strategy 1: Simple case - placeholder in single run
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return True

        # Strategy 2: Complex case - placeholder fragmented across runs
        # Rebuild full text and track run boundaries
        full_text = paragraph.text
        runs_info = []
        char_index = 0

        for run in paragraph.runs:
            run_length = len(run.text)
            runs_info.append({
                'run': run,
                'start': char_index,
                'end': char_index + run_length,
                'original_text': run.text,
                'font': run.font.name if run.font.name else None,
                'size': run.font.size,
                'bold': run.bold,
                'italic': run.italic
            })
            char_index += run_length

        # Find and replace in full text
        if old_text in full_text:
            new_full_text = full_text.replace(old_text, new_text)

            # Clear existing runs
            for _ in range(len(paragraph.runs)):
                paragraph._element.remove(paragraph.runs[0]._element)

            # Add single new run with complete text
            new_run = paragraph.add_run(new_full_text)

            # Try to preserve formatting from first run
            if runs_info:
                first_run_info = runs_info[0]
                if first_run_info['font']:
                    new_run.font.name = first_run_info['font']
                if first_run_info['size']:
                    new_run.font.size = first_run_info['size']
                new_run.bold = first_run_info['bold']
                new_run.italic = first_run_info['italic']

            return True

        return False

    def _fill_portada(self, context: Dict[str, Any]):
        """Fill cover page with report name - ROBUST VERSION"""
        nombre_tablero = context['nombre_tablero']

        # Method 1: Replace in all paragraphs using robust method
        for paragraph in self.doc.paragraphs:
            self._replace_text_in_paragraph(paragraph, '[Nombre del tablero]', nombre_tablero)
            self._replace_text_in_paragraph(paragraph, 'Nombre del tablero', nombre_tablero)

        # Method 2: Replace in headers/footers
        for section in self.doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    self._replace_text_in_paragraph(paragraph, '[Nombre del tablero]', nombre_tablero)
                    self._replace_text_in_paragraph(paragraph, 'Nombre del tablero', nombre_tablero)
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    self._replace_text_in_paragraph(paragraph, '[Nombre del tablero]', nombre_tablero)
                    self._replace_text_in_paragraph(paragraph, 'Nombre del tablero', nombre_tablero)

        # Method 3: Replace in text boxes (keep existing XML-based approach)
        try:
            # Get all shapes in the document
            for element in self.doc.element.body.iter():
                # Check for drawing elements
                if element.tag.endswith('}txbxContent'):
                    # This is a textbox content
                    for p_elem in element.iter():
                        if p_elem.tag.endswith('}t'):  # Text element
                            if p_elem.text:
                                p_elem.text = p_elem.text.replace('[Nombre del tablero]', nombre_tablero)
                                p_elem.text = p_elem.text.replace('Nombre del tablero', nombre_tablero)
        except Exception as e:
            self.logger.warning(f"Could not replace in shapes: {e}")

        self.logger.info(f"Filled portada with: {nombre_tablero}")

    def _fill_version_table(self, context: Dict[str, Any]):
        """Fill version history table"""
        if len(self.doc.tables) == 0:
            self.logger.warning("No tables found for version history")
            return

        table = self.doc.tables[0]

        if len(table.rows) >= 2:
            versiones = context['versiones']
            if versiones:
                v = versiones[0]
                cells = table.rows[1].cells

                cells[0].text = v['version']
                cells[1].text = v['fecha']
                cells[2].text = v['autor']
                cells[3].text = v['observaciones']

        self.logger.info("Filled version table")

    def _fill_objective(self, context: Dict[str, Any]):
        """Fill objective section"""
        objetivo = context['objetivo']
        self._replace_section_content('Objetivo', objetivo)

    def _fill_scope(self, context: Dict[str, Any]):
        """Fill scope section"""
        alcance = context['alcance']
        self._replace_section_content('Alcance', alcance)

    def _fill_users(self, context: Dict[str, Any]):
        """Fill users section"""
        usuarios = context['usuarios']
        self._replace_section_content('Usuarios', usuarios)

    def _fill_definitions(self, context: Dict[str, Any]):
        """Fill definitions section with measures"""
        medidas_por_categoria = context['medidas_por_categoria']

        if not medidas_por_categoria:
            self.logger.warning("No measures found to document")
            return

        # Find "Definiciones" heading
        heading_para = self._find_heading_paragraph('Definiciones')
        if not heading_para:
            self.logger.warning("Definiciones heading not found")
            return

        # Get the parent element (body)
        parent = heading_para._element.getparent()
        heading_elem = heading_para._element

        # Remove old content after heading (placeholders)
        self._remove_content_after_heading(heading_para)

        # Build content
        content_paras = []

        for categoria_data in medidas_por_categoria:
            categoria = categoria_data['categoria']
            medidas = categoria_data['medidas']

            # Category heading
            content_paras.append(('heading2', f"Categoría: {categoria}"))

            # Add each measure
            for medida in medidas:
                # Measure name
                content_paras.append(('bold', f"• {medida['nombre']}"))

                # Formula
                if medida['formula']:
                    content_paras.append(('normal', "  Fórmula DAX:"))
                    content_paras.append(('code', f"  {medida['formula']}"))

                # Format
                if medida['formato']:
                    content_paras.append(('normal', f"  Formato: {medida['formato']}"))

                # Description
                if medida['descripcion']:
                    content_paras.append(('normal', f"  Descripción: {medida['descripcion']}"))

                # Spacing
                content_paras.append(('normal', ''))

        # Insert content after heading
        self._insert_paragraphs_after(heading_para, content_paras)
        self.logger.info(f"Filled {len(medidas_por_categoria)} measure categories")

    def _fill_data_sources(self, context: Dict[str, Any]):
        """Fill data sources section"""
        tablas_origenes = context['tablas_origenes']

        if not tablas_origenes:
            self.logger.warning("No data sources found")
            return

        heading_para = self._find_heading_paragraph('Orígenes')
        if not heading_para:
            return

        self._remove_content_after_heading(heading_para)

        content_paras = []

        for tabla_info in tablas_origenes:
            # Table name
            content_paras.append(('heading3', f"Tabla: {tabla_info['tabla']}"))

            # Source type
            content_paras.append(('normal', f"Tipo de fuente: {tabla_info['tipo_fuente']}"))

            # Mode
            content_paras.append(('normal', f"Modo de conexión: {tabla_info['modo']}"))

            # Query M (if available)
            if tabla_info['query_m']:
                content_paras.append(('bold', "Power Query M:"))
                content_paras.append(('code', tabla_info['query_m']))

            # Spacing
            content_paras.append(('normal', ''))

        self._insert_paragraphs_after(heading_para, content_paras)
        self.logger.info(f"Filled {len(tablas_origenes)} data sources")

    def _fill_filters(self, context: Dict[str, Any]):
        """Fill filters section with REPORT-LEVEL filters"""
        heading_para = self._find_heading_paragraph('Filtros')
        if not heading_para:
            return

        self._remove_content_after_heading(heading_para)

        content_paras = []

        # Get new filter data
        filtros_reporte = context.get('filtros_reporte', [])
        filtros_pagina = context.get('filtros_pagina', [])
        slicers = context.get('slicers', [])

        # 1. Report-level filters (global)
        if filtros_reporte:
            content_paras.append(('heading2', f"Filtros a Nivel de Reporte ({len(filtros_reporte)})"))
            content_paras.append(('normal', 'Estos filtros se aplican a todas las páginas del reporte:'))
            content_paras.append(('normal', ''))
            for filtro in filtros_reporte:
                content_paras.append(('normal', f"• {filtro['field']} ({filtro['type']})"))
                if filtro.get('expression'):
                    content_paras.append(('normal', f"  └─ {filtro['expression']}"))
            content_paras.append(('normal', ''))

        # 2. Page-level filters
        if filtros_pagina:
            content_paras.append(('heading2', f"Filtros a Nivel de Página ({len(filtros_pagina)})"))
            content_paras.append(('normal', 'Estos filtros se aplican a páginas específicas:'))
            content_paras.append(('normal', ''))

            # Group by page
            by_page = {}
            for filtro in filtros_pagina:
                page = filtro['scope']
                if page not in by_page:
                    by_page[page] = []
                by_page[page].append(filtro)

            for page, filters in by_page.items():
                content_paras.append(('normal', f"📄 {page}:"))
                for filtro in filters:
                    content_paras.append(('normal', f"  • {filtro['field']} ({filtro['type']})"))
                    if filtro.get('expression'):
                        content_paras.append(('normal', f"    └─ {filtro['expression']}"))
                content_paras.append(('normal', ''))

        # 3. Slicers
        if slicers:
            content_paras.append(('heading2', f"Segmentaciones (Slicers) ({len(slicers)})"))
            content_paras.append(('normal', 'Controles interactivos de filtrado disponibles para el usuario:'))
            content_paras.append(('normal', ''))

            # Group by page
            by_page = {}
            for slicer in slicers:
                page = slicer['page']
                if page not in by_page:
                    by_page[page] = []
                by_page[page].append(slicer)

            for page, page_slicers in by_page.items():
                content_paras.append(('normal', f"📄 {page}:"))
                for slicer in page_slicers:
                    content_paras.append(('normal', f"  • {slicer['field']} (Tipo: {slicer['type']})"))
                content_paras.append(('normal', ''))

        # 4. Visual-level filters
        filtros_visual = context.get('filtros_visual', [])
        if filtros_visual:
            content_paras.append(('heading2', f"Filtros a Nivel de Visual ({len(filtros_visual)})"))
            content_paras.append(('normal', 'Estos filtros se aplican a gráficos específicos:'))
            content_paras.append(('normal', ''))

            # Group by scope (which includes visual type and page)
            by_scope = {}
            for filtro in filtros_visual:
                scope = filtro['scope']  # e.g., "Visual: cardVisual (Página Principal)"
                if scope not in by_scope:
                    by_scope[scope] = []
                by_scope[scope].append(filtro)

            for scope, filters in by_scope.items():
                content_paras.append(('normal', f"📊 {scope}:"))
                for filtro in filters:
                    content_paras.append(('normal', f"  • {filtro['field']} ({filtro['type']})"))
                    if filtro.get('expression'):
                        content_paras.append(('normal', f"    └─ {filtro['expression']}"))
                content_paras.append(('normal', ''))

        # 5. If no report filters at all, check for legacy Power Query/DAX filters
        if not filtros_reporte and not filtros_pagina and not slicers and not filtros_visual:
            # Fallback to old Power Query/DAX detection
            filtros_pq = context.get('filtros_powerquery', [])
            filtros_dax = context.get('filtros_dax', [])

            if filtros_pq or filtros_dax:
                # Has data model filters, show them
                if filtros_pq:
                    content_paras.append(('heading2', "Filtros en Power Query"))
                    for filtro in filtros_pq:
                        content_paras.append(('normal', f"• {filtro['tabla']}: {filtro['descripcion']}"))
                    content_paras.append(('normal', ''))

                if filtros_dax:
                    content_paras.append(('heading2', "Filtros DAX (Tablas Calculadas)"))
                    for filtro in filtros_dax:
                        content_paras.append(('normal', f"• {filtro['tabla']}: {filtro['descripcion']}"))
                    content_paras.append(('normal', ''))
            else:
                # No filters at all
                content_paras.append(('normal', "No se detectaron filtros en el reporte."))
                content_paras.append(('normal', ""))
                content_paras.append(('normal', "Nota: El modelo utiliza conexión DirectQuery sin transformaciones Power Query,"))
                content_paras.append(('normal', "y no se encontraron filtros aplicados a nivel de reporte, página o visual."))

        self._insert_paragraphs_after(heading_para, content_paras)
        self.logger.info(f"Filled filters section: {len(filtros_reporte)} report, {len(filtros_pagina)} page, {len(slicers)} slicers, {len(filtros_visual)} visual")

    def _fill_er_model(self, context: Dict[str, Any], er_diagram_path: Optional[str],
                      er_image_path: Optional[str] = None):
        """Fill ER model section with user-uploaded or auto-generated diagram"""
        relaciones = context['relaciones']
        stats = context['estadisticas_modelo']

        heading_para = self._find_heading_paragraph('Modelo ER')
        if not heading_para:
            return

        self._remove_content_after_heading(heading_para)

        content_paras = []

        # Prioritize user-uploaded ER image over auto-generated
        er_image_to_use = None
        image_source = ""

        if er_image_path and Path(er_image_path).exists():
            er_image_to_use = er_image_path
            image_source = " (imagen proporcionada por el usuario)"
            self.logger.info(f"Using user-uploaded ER image: {er_image_path}")
        elif er_diagram_path and Path(er_diagram_path).exists():
            er_image_to_use = er_diagram_path
            image_source = " (generada automáticamente)"
            self.logger.info(f"Using auto-generated ER diagram: {er_diagram_path}")

        # Add ER diagram if available
        if er_image_to_use:
            content_paras.append(('heading2', f"Diagrama de Relaciones{image_source}"))
            content_paras.append(('image', er_image_to_use))

        # Statistics
        content_paras.append(('heading2', "Estadísticas del Modelo"))
        stats_text = f"""• Total de tablas: {stats['total_tablas']}
• Total de relaciones: {stats['total_relaciones']}
• Relaciones activas: {stats['relaciones_activas']}
• Relaciones bidireccionales: {stats['relaciones_bidireccionales']}
• Total de medidas: {stats['total_medidas']}
• Auto Date/Time: {stats.get('auto_datetime', 'No detectado')}"""
        content_paras.append(('normal', stats_text))

        # Add hierarchies and columns here (from appendix)
        jerarquias = context.get('jerarquias', [])
        columnas = context.get('columnas_por_tabla', [])
        calc_cols = context.get('columnas_calculadas', [])

        if jerarquias:
            content_paras.append(('heading2', "Jerarquías Definidas"))
            for h in jerarquias:
                niveles = ' → '.join(h['niveles'])
                content_paras.append(('normal', f"• {h['nombre']} ({h['tabla']}): {niveles}"))

        # Columnas Calculadas
        if calc_cols:
            content_paras.append(('heading2', f"Columnas Calculadas ({len(calc_cols)} en total)"))
            # Group by table
            calc_by_table = {}
            for col in calc_cols:
                tabla = col.get('tabla', 'Desconocida')
                if tabla not in calc_by_table:
                    calc_by_table[tabla] = []
                calc_by_table[tabla].append(col)

            for tabla, cols in list(calc_by_table.items())[:10]:  # Max 10 tables
                content_paras.append(('bold', f"Tabla: {tabla}"))
                for col in cols[:5]:  # Max 5 columns per table
                    expr = col.get('expression', 'N/A')
                    if len(expr) > 100:
                        expr = expr[:100] + "..."
                    content_paras.append(('normal', f"  • {col.get('nombre', 'N/A')} = {expr}"))
                content_paras.append(('normal', ''))

        # Show first 10 tables with columns
        if columnas:
            content_paras.append(('heading2', "Detalle de Tablas y Columnas"))
            for tabla_cols in columnas[:10]:
                # Skip LocalDateTable
                if 'LocalDateTable' in tabla_cols['tabla']:
                    continue
                content_paras.append(('bold', f"Tabla: {tabla_cols['tabla']}"))
                cols_summary = []
                for col in tabla_cols['columnas'][:15]:  # Max 15 columns per table
                    tipo_col = " (calculada)" if col['calculada'] else ""
                    cols_summary.append(f"  • {col['nombre']} ({col['tipo']}){tipo_col}")
                content_paras.append(('normal', '\n'.join(cols_summary)))
                content_paras.append(('normal', ''))

        # Add relationships table heading before inserting
        if relaciones:
            content_paras.append(('heading2', f"Tabla de Relaciones ({len(relaciones)} relaciones)"))
            content_paras.append(('normal', ''))

        self._insert_paragraphs_after(heading_para, content_paras)

        # Relationships table - insert INLINE in Modelo ER section
        if relaciones:
            # Find the last paragraph we inserted
            parent = heading_para._element.getparent()
            heading_idx = list(parent).index(heading_para._element)

            # Find where to insert (after all our content)
            insert_position = heading_idx + len(content_paras) + 1

            # Create table
            table = self.doc.add_table(rows=1, cols=5)
            table_elem = table._element

            # Move table to correct position (not at end of document)
            parent.remove(table_elem)
            parent.insert(insert_position, table_elem)

            try:
                table.style = 'Light Grid Accent 1'
            except KeyError:
                try:
                    table.style = 'Table Grid'
                except KeyError:
                    pass

            # Headers
            headers = ['Desde', 'Hacia', 'Cardinalidad', 'Dirección', 'Estado']
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Rows
            for rel in relaciones:
                row_cells = table.add_row().cells
                row_cells[0].text = rel['desde']
                row_cells[1].text = rel['hacia']
                row_cells[2].text = rel['cardinalidad']
                row_cells[3].text = rel['direccion']
                row_cells[4].text = rel['estado']

        self.logger.info(f"Filled ER model with {len(relaciones)} relationships")

    def _fill_visualizations_section(self, context: Dict[str, Any],
                                     visualization_images: List[str]):
        """Fill visualizations section with user-uploaded images"""
        if not visualization_images:
            self.logger.info("No visualization images provided, skipping section")
            return

        # Find ER Model section to insert after it
        er_heading = self._find_heading_paragraph('Modelo ER')
        if not er_heading:
            self.logger.warning("Could not find 'Modelo ER' heading to insert visualizations after")
            return

        # Find the position to insert the new section
        # We need to insert after all the ER model content
        parent = er_heading._element.getparent()
        er_heading_idx = list(parent).index(er_heading._element)

        # Find the next Heading 1 (to know where ER section ends)
        # Only match Heading 1 to avoid matching sub-headings (Heading 2/3) within the ER section
        next_heading_idx = None
        for idx, elem in enumerate(list(parent)[er_heading_idx + 1:], start=er_heading_idx + 1):
            if hasattr(elem, 'style'):
                try:
                    para = next((p for p in self.doc.paragraphs if p._element == elem), None)
                    if para and para.style.name == 'Heading 1':
                        next_heading_idx = idx
                        break
                except:
                    pass

        # Insert position: right before the next heading or at the end
        if next_heading_idx:
            insert_position = next_heading_idx
        else:
            insert_position = len(list(parent))

        # Create new heading for visualizations
        heading_para = self.doc.add_paragraph()
        heading_para.style = 'Heading 1'
        heading_para.add_run("Visualizaciones del Reporte").bold = True

        # Move heading to correct position
        heading_elem = heading_para._element
        parent.remove(heading_elem)
        parent.insert(insert_position, heading_elem)
        insert_position += 1

        # Add intro paragraph
        intro_para = self.doc.add_paragraph()
        intro_para.add_run("A continuación se presentan las principales visualizaciones y dashboards del reporte:")
        intro_elem = intro_para._element
        parent.remove(intro_elem)
        parent.insert(insert_position, intro_elem)
        insert_position += 1

        # Add spacing
        space_para = self.doc.add_paragraph()
        space_elem = space_para._element
        parent.remove(space_elem)
        parent.insert(insert_position, space_elem)
        insert_position += 1

        # Add each visualization image
        for idx, img_path in enumerate(visualization_images, start=1):
            if not Path(img_path).exists():
                self.logger.warning(f"Visualization image not found: {img_path}")
                continue

            # Add figure caption
            caption_para = self.doc.add_paragraph()
            caption_para.add_run(f"Figura {idx}: Visualización del Reporte").bold = True
            caption_elem = caption_para._element
            parent.remove(caption_elem)
            parent.insert(insert_position, caption_elem)
            insert_position += 1

            # Add image
            img_para = self.doc.add_paragraph()
            run = img_para.add_run()
            try:
                run.add_picture(img_path, width=Inches(6.0))
                img_elem = img_para._element
                parent.remove(img_elem)
                parent.insert(insert_position, img_elem)
                insert_position += 1
                self.logger.info(f"Added visualization image {idx}: {img_path}")
            except Exception as e:
                self.logger.error(f"Error adding visualization image {img_path}: {e}")
                # Remove the paragraph if image failed
                parent.remove(img_para._element)

            # Add spacing between images
            space_para = self.doc.add_paragraph()
            space_elem = space_para._element
            parent.remove(space_elem)
            parent.insert(insert_position, space_elem)
            insert_position += 1

        self.logger.info(f"Filled visualizations section with {len(visualization_images)} images")

    def _fill_rls(self, context: Dict[str, Any]):
        """Fill RLS/Security section"""
        roles = context['roles_rls']
        tiene_rls = context['tiene_rls']

        heading_para = self._find_heading_paragraph('Consideraciones técnicas')
        if not heading_para:
            # Try alternative heading
            heading_para = self._find_heading_paragraph('Consideraciones')

        if not heading_para:
            return

        self._remove_content_after_heading(heading_para)

        content_paras = []

        # Refresh frequency
        content_paras.append(('normal', f"Frecuencia de actualización: {context['frecuencia_actualizacion']}"))

        # Connection mode
        content_paras.append(('normal', f"Modo de conexión principal: {context['modo_conexion']}"))

        # Spacing
        content_paras.append(('normal', ''))

        # RLS section
        content_paras.append(('heading2', "Seguridad a Nivel de Fila (RLS)"))

        if not tiene_rls:
            content_paras.append(('normal', "No se detectaron roles de seguridad configurados."))
        else:
            for role in roles:
                # Role name
                content_paras.append(('bold', f"Rol: {role['nombre']}"))

                # Permissions
                for permiso in role['permisos']:
                    content_paras.append(('normal', f"  • Tabla: {permiso['tabla']}"))
                    if permiso['filtro']:
                        content_paras.append(('code', f"    Filtro: {permiso['filtro']}"))

                # Spacing
                content_paras.append(('normal', ''))

        self._insert_paragraphs_after(heading_para, content_paras)
        self.logger.info("Filled RLS section")

    def _fill_appendix(self, context: Dict[str, Any]):
        """Fill appendix section with Power Query M expressions"""
        heading_para = self._find_heading_paragraph('Anexo')
        if not heading_para:
            return

        self._remove_content_after_heading(heading_para)

        content_paras = []
        power_queries = context.get('power_queries', [])

        if power_queries:
            content_paras.append(('heading2', 'Código Power Query (M)'))
            content_paras.append(('normal',
                f"Las siguientes {len(power_queries)} consultas de Power Query definen "
                "las transformaciones de datos y conexiones a fuentes de datos del modelo."
            ))

            for pq in power_queries:
                tabla = pq['tabla']
                tipo = pq.get('tipo_fuente', 'Import')
                content_paras.append(('heading3', f"{tabla} ({tipo})"))
                content_paras.append(('code', pq['expresion']))
        else:
            content_paras.append(('normal',
                "Consultar sección 'Modelo ER' para el detalle completo del modelo de datos."
            ))

        self._insert_paragraphs_after(heading_para, content_paras)
        self.logger.info(f"Filled appendix with {len(power_queries)} Power Query expressions")

    # === HELPER METHODS ===

    def _replace_section_content(self, heading_text: str, content: str):
        """Replace content after a heading with new content"""
        heading_para = self._find_heading_paragraph(heading_text)
        if not heading_para:
            self.logger.warning(f"Heading '{heading_text}' not found")
            return

        self._remove_content_after_heading(heading_para)
        self._insert_paragraphs_after(heading_para, [('normal', content)])

    def _insert_paragraphs_after(self, heading_para, content_paras: List[tuple]):
        """
        Insert paragraphs after a heading

        Args:
            heading_para: The heading paragraph
            content_paras: List of tuples (style, text) where style is:
                - 'normal': Normal paragraph
                - 'bold': Bold text
                - 'heading2': Heading 2
                - 'heading3': Heading 3
                - 'code': Code block
                - 'image': Image path
        """
        parent = heading_para._element.getparent()
        heading_elem = heading_para._element

        # Find position after heading
        position = parent.index(heading_elem) + 1

        for style, text in content_paras:
            para = self.doc.add_paragraph()
            para_elem = para._element

            # Move paragraph to correct position
            parent.remove(para_elem)
            parent.insert(position, para_elem)
            position += 1

            # Apply style and content
            if style == 'heading2':
                para.style = 'Heading 2'
                para.add_run(text).bold = True
            elif style == 'heading3':
                para.style = 'Heading 3'
                para.add_run(text).bold = True
            elif style == 'bold':
                para.add_run(text).bold = True
            elif style == 'code':
                run = para.add_run(text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            elif style == 'image':
                if Path(text).exists():
                    run = para.add_run()
                    run.add_picture(text, width=Inches(6.0))
            else:  # normal
                para.add_run(text)

    def _remove_content_after_heading(self, heading_para):
        """Remove placeholder content after heading"""
        parent = heading_para._element.getparent()
        heading_elem = heading_para._element

        # Find position after heading
        position = parent.index(heading_elem) + 1

        # Check if next paragraph is a placeholder
        if position < len(parent):
            next_elem = parent[position]
            if next_elem.tag.endswith('p'):  # It's a paragraph
                # Create temporary paragraph to access text
                temp_para = None
                for para in self.doc.paragraphs:
                    if para._element == next_elem:
                        temp_para = para
                        break

                if temp_para and temp_para.text.strip():
                    # Check if it's a placeholder (starts with [ and ends with ])
                    text = temp_para.text.strip()
                    if text.startswith('[') and text.endswith(']'):
                        parent.remove(next_elem)

    def _find_heading_paragraph(self, heading_text: str):
        """Find paragraph containing heading text"""
        for para in self.doc.paragraphs:
            if heading_text.lower() in para.text.lower():
                style = para.style.name if para.style else ''
                if 'Heading' in style or 'Título' in style or 'Ttulo' in style:
                    return para
        return None

    def _replace_text_in_paragraph(self, paragraph, old_text: str, new_text: str):
        """Replace text in paragraph while preserving formatting"""
        if old_text in paragraph.text:
            # Simple replacement
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    run.bold = True

    def _get_output_path(self, context: Dict[str, Any]) -> Path:
        """Generate output path"""
        output_dir = Path(__file__).resolve().parent.parent.parent.parent / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_name = context['nombre_tablero'].replace(" ", "_")
        report_name = "".join(c for c in report_name if c.isalnum() or c in ('_', '-'))

        filename = f"PowerBI_Doc_{report_name}_{timestamp}.docx"
        return output_dir / filename

    def _progress(self, callback, step: int, message: str):
        """Call progress callback if provided"""
        if callback:
            callback(step, message)
        self.logger.info(f"[{step}%] {message}")
