"""
DAX Measures Section Generator

Genera documentación completa de medidas DAX con expresiones COMPLETAS (SIN TRUNCAR).
"""

from docx import Document
from typing import Any
import logging

from ..utils.docx_helpers import DocxHelpers
from ..dax_description_generator import DAXDescriptionGenerator

logger = logging.getLogger(__name__)


class DAXMeasuresSectionGenerator:
    """Genera sección de documentación de medidas DAX"""

    def __init__(self, doc: Document, metadata: Any, template_handler: Any):
        """
        Initialize DAX measures section generator

        Args:
            doc: Document object
            metadata: ReportMetadata object
            template_handler: TemplateHandler object
        """
        self.doc = doc
        self.metadata = metadata
        self.template_handler = template_handler
        self.description_generator = DAXDescriptionGenerator()

    def generate(self) -> None:
        """Generate DAX measures section"""
        logger.info("Generating DAX measures section")

        DocxHelpers.add_heading(self.doc, 'Medidas DAX', level=1)

        if not hasattr(self.metadata, 'data_model') or not self.metadata.data_model:
            DocxHelpers.add_paragraph(self.doc, "No hay información del modelo de datos disponible.")
            return

        model = self.metadata.data_model
        measures = model.measures if hasattr(model, 'measures') else []

        if not measures:
            DocxHelpers.add_paragraph(self.doc, "No hay medidas DAX definidas en el modelo de datos.")
            return

        # Overview
        self._add_overview(measures)

        # Complexity summary
        self._add_complexity_summary(measures)

        # Measures grouped by complexity
        self._add_measures_by_complexity(measures)

        # Most complex measures
        self._add_most_complex_measures(measures)

        logger.info("DAX measures section generated successfully")

    def _add_overview(self, measures: list) -> None:
        """Add overview paragraph"""
        overview_text = (
            f"El modelo de datos contiene {len(measures)} medidas DAX. "
            "Estas medidas realizan cálculos y agregaciones a través del modelo de datos. "
            "Las siguientes secciones proporcionan documentación completa incluyendo expresiones, "
            "análisis de complejidad y uso de funciones."
        )
        DocxHelpers.add_paragraph(self.doc, overview_text)

    def _add_complexity_summary(self, measures: list) -> None:
        """Add complexity summary"""
        DocxHelpers.add_heading(self.doc, 'Resumen de Complejidad', level=2)

        # Count by complexity
        complexity_counts = {
            'LOW': 0,
            'MEDIUM': 0,
            'HIGH': 0,
            'VERY_HIGH': 0
        }

        for measure in measures:
            complexity = getattr(measure, 'complexity', 'UNKNOWN')
            if complexity in complexity_counts:
                complexity_counts[complexity] += 1

        # Create summary table
        data = [
            ['Complejidad Baja', str(complexity_counts['LOW']), 'Cálculos simples, agregaciones básicas'],
            ['Complejidad Media', str(complexity_counts['MEDIUM']), 'Múltiples funciones, cierta lógica'],
            ['Complejidad Alta', str(complexity_counts['HIGH']), 'Lógica compleja, iteradores, inteligencia de tiempo'],
            ['Complejidad Muy Alta', str(complexity_counts['VERY_HIGH']), 'Muy complejo, iteradores anidados, múltiples contextos']
        ]
        headers = ['Nivel de Complejidad', 'Cantidad', 'Descripción']
        DocxHelpers.add_table(self.doc, data, headers, style='Table Grid')

    def _add_measures_by_complexity(self, measures: list) -> None:
        """Add all measures grouped by complexity"""
        DocxHelpers.add_heading(self.doc, 'Medidas por Complejidad', level=2)

        # Group measures by complexity
        grouped = {
            'LOW': [],
            'MEDIUM': [],
            'HIGH': [],
            'VERY_HIGH': []
        }

        for measure in measures:
            complexity = getattr(measure, 'complexity', 'UNKNOWN')
            if complexity in grouped:
                grouped[complexity].append(measure)

        # Document each group
        complexity_labels = {
            'LOW': 'Baja',
            'MEDIUM': 'Media',
            'HIGH': 'Alta',
            'VERY_HIGH': 'Muy Alta'
        }

        for complexity in ['LOW', 'MEDIUM', 'HIGH', 'VERY_HIGH']:
            complexity_measures = grouped[complexity]

            if not complexity_measures:
                continue

            # Complexity heading
            complexity_title = f"Complejidad {complexity_labels[complexity]}"
            DocxHelpers.add_heading(self.doc, complexity_title, level=3)

            # Color badge based on complexity
            badge_para = self.doc.add_paragraph()
            if complexity == 'LOW':
                DocxHelpers.add_colored_badge(badge_para, f"{len(complexity_measures)} medidas", 'green')
            elif complexity == 'MEDIUM':
                DocxHelpers.add_colored_badge(badge_para, f"{len(complexity_measures)} medidas", 'yellow')
            elif complexity == 'HIGH':
                DocxHelpers.add_colored_badge(badge_para, f"{len(complexity_measures)} medidas", 'orange')
            else:  # VERY_HIGH
                DocxHelpers.add_colored_badge(badge_para, f"{len(complexity_measures)} medidas", 'red')

            # Document each measure
            for measure in complexity_measures:
                self._document_measure(measure, complexity)

    def _document_measure(self, measure: Any, complexity: str) -> None:
        """
        Document a single measure

        Args:
            measure: Measure object
            complexity: Complexity level
        """
        # Measure name
        measure_name = getattr(measure, 'name', 'Desconocido')
        DocxHelpers.add_heading(self.doc, measure_name, level=4)

        # === BUSINESS DESCRIPTION ===
        business_desc = self.description_generator.generate_description(
            measure_name=measure_name,
            expression=getattr(measure, 'expression', ''),
            table_name=getattr(measure, 'table', ''),
            existing_description=getattr(measure, 'description', None)
        )

        if business_desc:
            DocxHelpers.add_heading(self.doc, 'Descripción', level=5)
            DocxHelpers.add_paragraph(self.doc, business_desc, style='Body Text')

        # Metadata
        metadata_items = []

        if hasattr(measure, 'table'):
            metadata_items.append(f"Tabla: {measure.table}")

        complexity_labels = {
            'LOW': 'Baja',
            'MEDIUM': 'Media',
            'HIGH': 'Alta',
            'VERY_HIGH': 'Muy Alta'
        }
        complexity_label = complexity_labels.get(complexity, complexity)
        metadata_items.append(f"Complejidad: {complexity_label}")

        if hasattr(measure, 'expression'):
            expression_length = len(measure.expression)
            metadata_items.append(f"Longitud de Expresión: {expression_length} caracteres")

        if metadata_items:
            DocxHelpers.add_bulleted_list(self.doc, metadata_items)

        # Expression - COMPLETA, SIN TRUNCAR!
        if hasattr(measure, 'expression') and measure.expression:
            DocxHelpers.add_heading(self.doc, 'Expresión', level=5)
            # Agregar la expresión COMPLETA sin ningún truncamiento
            DocxHelpers.add_code_block(self.doc, measure.expression, language='dax')

        # Key characteristics
        if hasattr(measure, 'analysis'):
            analysis = measure.analysis
            characteristics = []

            if hasattr(analysis, 'uses_time_intelligence'):
                characteristics.append(f"Usa Inteligencia de Tiempo: {'Sí' if analysis.uses_time_intelligence else 'No'}")

            if hasattr(analysis, 'uses_iterators'):
                characteristics.append(f"Usa Iteradores: {'Sí' if analysis.uses_iterators else 'No'}")

            if hasattr(analysis, 'uses_context_transition'):
                characteristics.append(f"Usa Transición de Contexto: {'Sí' if analysis.uses_context_transition else 'No'}")

            if hasattr(analysis, 'function_usage') and analysis.function_usage:
                func_list = ', '.join([f"{func}: {count}" for func, count in analysis.function_usage.items()])
                characteristics.append(f"Funciones Usadas: {func_list}")

            if characteristics:
                DocxHelpers.add_heading(self.doc, 'Características', level=5)
                DocxHelpers.add_bulleted_list(self.doc, characteristics)

        # Add spacing
        self.doc.add_paragraph()

    def _add_most_complex_measures(self, measures: list) -> None:
        """Add detailed analysis of most complex measures"""
        # Filter for high and very high complexity
        complex_measures = [
            m for m in measures
            if hasattr(m, 'complexity') and m.complexity in ['HIGH', 'VERY_HIGH']
        ]

        if not complex_measures:
            return

        # Sort by complexity score if available
        if hasattr(complex_measures[0], 'complexity_score'):
            complex_measures.sort(key=lambda m: getattr(m, 'complexity_score', 0), reverse=True)

        # Take top 10
        top_measures = complex_measures[:10]

        DocxHelpers.add_heading(self.doc, 'Medidas Más Complejas (Top 10)', level=2)

        DocxHelpers.add_emphasis_box(
            self.doc,
            f"Las siguientes {len(top_measures)} medidas tienen la mayor complejidad. "
            "Revisar estas para oportunidades de optimización o para asegurar que funcionen según lo esperado.",
            title="⚠️ Atención Requerida"
        )

        # Document each in detail
        for idx, measure in enumerate(top_measures, 1):
            measure_name = getattr(measure, 'name', 'Desconocido')
            DocxHelpers.add_heading(self.doc, f"{idx}. {measure_name}", level=3)

            # Full documentation
            self._document_measure(measure, getattr(measure, 'complexity', 'UNKNOWN'))
