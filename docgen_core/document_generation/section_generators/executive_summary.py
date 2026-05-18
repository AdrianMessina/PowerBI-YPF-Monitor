"""
Executive Summary Generator

Genera resumen ejecutivo con estadísticas clave e insights.
"""

from docx import Document
from typing import Any
import logging

from ..utils.docx_helpers import DocxHelpers

logger = logging.getLogger(__name__)


class ExecutiveSummaryGenerator:
    """Genera sección de resumen ejecutivo"""

    def __init__(self, doc: Document, metadata: Any, validation_report: Any,
                 template_handler: Any):
        """
        Initialize executive summary generator

        Args:
            doc: Document object
            metadata: ReportMetadata object
            validation_report: ValidationReport object
            template_handler: TemplateHandler object
        """
        self.doc = doc
        self.metadata = metadata
        self.validation_report = validation_report
        self.template_handler = template_handler

    def generate(self) -> None:
        """Generate executive summary"""
        logger.info("Generating executive summary")

        # Add section heading
        DocxHelpers.add_heading(self.doc, 'Resumen Ejecutivo', level=1)

        # Add overview
        self._add_overview()

        # Add key metrics
        self._add_key_metrics()

        # Add quality score
        self._add_quality_score()

        # Add quick insights
        self._add_quick_insights()

        # Page break after summary
        DocxHelpers.add_page_break(self.doc)

        logger.info("Executive summary generated successfully")

    def _add_overview(self) -> None:
        """Add report overview"""
        DocxHelpers.add_heading(self.doc, 'Descripción General', level=2)

        # Report name and description
        overview_text = f"Este documento proporciona documentación técnica completa del reporte de Power BI '{self.metadata.report_name}'."

        if hasattr(self.metadata, 'description') and self.metadata.description:
            overview_text += f" {self.metadata.description}"

        DocxHelpers.add_paragraph(self.doc, overview_text)

        # Scope
        scope_text = (
            "La documentación incluye información detallada sobre el modelo de datos, "
            "relaciones, medidas DAX, configuración de seguridad, visualizaciones, "
            "y resultados de validación de calidad."
        )
        DocxHelpers.add_paragraph(self.doc, scope_text)

    def _add_key_metrics(self) -> None:
        """Add key metrics table"""
        DocxHelpers.add_heading(self.doc, 'Métricas Clave', level=2)

        # Gather statistics
        stats = self._gather_statistics()

        # Create table
        data = [[label, str(value)] for label, value in stats.items()]
        headers = ['Métrica', 'Valor']

        DocxHelpers.add_table(self.doc, data, headers, style='Table Grid')

    def _gather_statistics(self) -> dict:
        """
        Gather report statistics

        Returns:
            Dictionary of statistics
        """
        stats = {}

        # Data model stats
        if hasattr(self.metadata, 'data_model') and self.metadata.data_model:
            model = self.metadata.data_model
            tables = model.tables if hasattr(model, 'tables') else []

            # Count tables by type
            fact_tables = sum(1 for t in tables if hasattr(t, 'table_type') and t.table_type == 'Fact')
            dim_tables = sum(1 for t in tables if hasattr(t, 'table_type') and t.table_type == 'Dimension')

            stats['Total de Tablas'] = len(tables)
            if fact_tables > 0:
                stats['  - Tablas de Hechos'] = fact_tables
            if dim_tables > 0:
                stats['  - Tablas de Dimensiones'] = dim_tables

            # Count columns
            total_columns = sum(len(t.columns) if hasattr(t, 'columns') else 0 for t in tables)
            stats['Total de Columnas'] = total_columns

            # Count calculated columns
            calc_columns = sum(
                sum(1 for c in t.columns if hasattr(c, 'is_calculated') and c.is_calculated)
                if hasattr(t, 'columns') else 0
                for t in tables
            )
            if calc_columns > 0:
                stats['  - Columnas Calculadas'] = calc_columns

            # Relationships
            if hasattr(model, 'relationships'):
                relationships = model.relationships
                stats['Total de Relaciones'] = len(relationships)

                # Count bidirectional
                bidirectional = sum(
                    1 for r in relationships
                    if hasattr(r, 'cross_filter_direction') and r.cross_filter_direction == 'Both'
                )
                if bidirectional > 0:
                    stats['  - Bidireccionales'] = bidirectional

                # Count many-to-many
                many_to_many = sum(
                    1 for r in relationships
                    if hasattr(r, 'from_cardinality') and hasattr(r, 'to_cardinality')
                    and r.from_cardinality == 'Many' and r.to_cardinality == 'Many'
                )
                if many_to_many > 0:
                    stats['  - Muchos a Muchos'] = many_to_many

            # DAX measures
            if hasattr(model, 'measures'):
                measures = model.measures
                stats['Total de Medidas DAX'] = len(measures)

                # Count by complexity if available
                if hasattr(self.metadata, 'dax_analysis') and self.metadata.dax_analysis:
                    complexity_counts = {}
                    for measure in measures:
                        if hasattr(measure, 'complexity'):
                            complexity = measure.complexity
                            complexity_counts[complexity] = complexity_counts.get(complexity, 0) + 1

                    complexity_map = {
                        'LOW': 'Baja',
                        'MEDIUM': 'Media',
                        'HIGH': 'Alta',
                        'VERY_HIGH': 'Muy Alta'
                    }
                    for complexity in ['LOW', 'MEDIUM', 'HIGH', 'VERY_HIGH']:
                        if complexity in complexity_counts:
                            stats[f'  - Complejidad {complexity_map[complexity]}'] = complexity_counts[complexity]

        # Security stats
        if hasattr(self.metadata, 'security') and self.metadata.security:
            security = self.metadata.security
            if hasattr(security, 'roles'):
                stats['Roles RLS'] = len(security.roles)

        # Visualization stats
        if hasattr(self.metadata, 'report_layout') and self.metadata.report_layout:
            layout = self.metadata.report_layout
            if hasattr(layout, 'pages'):
                stats['Páginas del Reporte'] = len(layout.pages)

                # Count visuals
                total_visuals = sum(
                    len(page.visuals) if hasattr(page, 'visuals') else 0
                    for page in layout.pages
                )
                stats['Total de Visualizaciones'] = total_visuals

        return stats

    def _add_quality_score(self) -> None:
        """Add quality score section"""
        if not self.validation_report:
            return

        DocxHelpers.add_heading(self.doc, 'Puntuación de Calidad', level=2)

        if hasattr(self.validation_report, 'quality_score'):
            score = self.validation_report.quality_score

            # Add score paragraph
            score_para = self.doc.add_paragraph()
            score_run = score_para.add_run(f"Puntuación General de Calidad: {score}/100")
            score_run.bold = True
            score_run.font.size = 14

            # Add interpretation
            if score >= 90:
                interpretation = "Excelente - El reporte sigue las mejores prácticas con problemas mínimos."
                color = 'green'
            elif score >= 75:
                interpretation = "Bueno - El reporte está bien estructurado con algunas mejoras menores necesarias."
                color = 'yellow'
            elif score >= 60:
                interpretation = "Regular - Varios problemas requieren atención para mejorar la calidad del reporte."
                color = 'orange'
            else:
                interpretation = "Necesita Mejoras - Múltiples problemas críticos deben ser atendidos."
                color = 'red'

            DocxHelpers.add_paragraph(self.doc, interpretation)

            # Add issue summary if available
            if hasattr(self.validation_report, 'issues'):
                issues = self.validation_report.issues
                critical = sum(1 for i in issues if hasattr(i, 'severity') and i.severity == 'CRITICAL')
                errors = sum(1 for i in issues if hasattr(i, 'severity') and i.severity == 'ERROR')
                warnings = sum(1 for i in issues if hasattr(i, 'severity') and i.severity == 'WARNING')

                summary_text = f"Problemas encontrados: {critical} críticos, {errors} errores, {warnings} advertencias"
                DocxHelpers.add_paragraph(self.doc, summary_text)

    def _add_quick_insights(self) -> None:
        """Add quick insights"""
        DocxHelpers.add_heading(self.doc, 'Insights Rápidos', level=2)

        insights = []

        # Most complex measures
        if hasattr(self.metadata, 'data_model') and self.metadata.data_model:
            model = self.metadata.data_model
            if hasattr(model, 'measures'):
                # Find most complex measures
                complex_measures = [
                    m for m in model.measures
                    if hasattr(m, 'complexity') and m.complexity in ['HIGH', 'VERY_HIGH']
                ]
                if complex_measures:
                    insights.append(
                        f"Se encontraron {len(complex_measures)} medidas DAX de alta complejidad que pueden requerir optimización"
                    )

        # Critical issues
        if self.validation_report and hasattr(self.validation_report, 'issues'):
            critical_issues = [
                i for i in self.validation_report.issues
                if hasattr(i, 'severity') and i.severity == 'CRITICAL'
            ]
            if critical_issues:
                insights.append(
                    f"{len(critical_issues)} problemas críticos requieren atención inmediata"
                )

        # Bidirectional relationships
        if hasattr(self.metadata, 'data_model') and self.metadata.data_model:
            model = self.metadata.data_model
            if hasattr(model, 'relationships'):
                bidirectional = [
                    r for r in model.relationships
                    if hasattr(r, 'cross_filter_direction') and r.cross_filter_direction == 'Both'
                ]
                if bidirectional:
                    insights.append(
                        f"{len(bidirectional)} relaciones bidireccionales detectadas - revisar impacto en el rendimiento"
                    )

        # RLS configuration
        if hasattr(self.metadata, 'security') and self.metadata.security:
            security = self.metadata.security
            if hasattr(security, 'roles') and security.roles:
                insights.append(
                    f"Seguridad a nivel de fila configurada con {len(security.roles)} roles"
                )

        # Add insights as bulleted list
        if insights:
            DocxHelpers.add_bulleted_list(self.doc, insights)
        else:
            DocxHelpers.add_paragraph(
                self.doc,
                "No se detectaron insights significativos. El reporte parece estar bien estructurado."
            )
