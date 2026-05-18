"""
Data Model Section Generator

Genera documentación completa del modelo de datos con diagrama ER y detalles de tablas.
"""

from docx import Document
from typing import Any
import logging

from ..utils.docx_helpers import DocxHelpers
from ..utils.er_diagram_embedder import ERDiagramEmbedder

logger = logging.getLogger(__name__)


class DataModelSectionGenerator:
    """Genera sección de documentación del modelo de datos"""

    def __init__(self, doc: Document, metadata: Any, template_handler: Any,
                 er_diagram_generator: Any = None):
        """
        Initialize data model section generator

        Args:
            doc: Document object
            metadata: ReportMetadata object
            template_handler: TemplateHandler object
            er_diagram_generator: Optional ERDiagramGenerator object
        """
        self.doc = doc
        self.metadata = metadata
        self.template_handler = template_handler
        self.er_diagram_generator = er_diagram_generator

    def generate(self) -> None:
        """Generate data model section"""
        logger.info("Generating data model section")

        # Add section heading
        DocxHelpers.add_heading(self.doc, 'Modelo de Datos', level=1)

        # Add overview
        self._add_overview()

        # Add ER diagram
        self._add_er_diagram()

        # Add table summary
        self._add_table_summary()

        # Add detailed table documentation
        self._add_detailed_tables()

        logger.info("Data model section generated successfully")

    def _add_overview(self) -> None:
        """Add data model overview"""
        if not hasattr(self.metadata, 'data_model') or not self.metadata.data_model:
            DocxHelpers.add_paragraph(
                self.doc,
                "No hay información del modelo de datos disponible."
            )
            return

        model = self.metadata.data_model
        tables = model.tables if hasattr(model, 'tables') else []

        overview_text = (
            f"El modelo de datos contiene {len(tables)} tablas organizadas en un esquema estrella o copo de nieve. "
            "Las siguientes secciones proporcionan documentación detallada de cada tabla incluyendo columnas, "
            "tipos de datos, relaciones y campos calculados."
        )

        DocxHelpers.add_paragraph(self.doc, overview_text)

    def _add_er_diagram(self) -> None:
        """Add ER diagram"""
        if not self.er_diagram_generator:
            logger.warning("No ER diagram generator available")
            return

        DocxHelpers.add_heading(self.doc, 'Diagrama Entidad-Relación', level=2)

        # Embed diagram
        embedder = ERDiagramEmbedder(self.doc)
        success = embedder.embed_diagram(
            self.er_diagram_generator,
            width_inches=6.5,
            caption="Figura: Diagrama Entidad-Relación del Modelo de Datos"
        )

        if not success:
            DocxHelpers.add_paragraph(
                self.doc,
                "No se pudo generar el diagrama ER. Por favor asegúrese de tener kaleido instalado para exportar imágenes estáticas."
            )

    def _add_table_summary(self) -> None:
        """Add table summary"""
        if not hasattr(self.metadata, 'data_model') or not self.metadata.data_model:
            return

        model = self.metadata.data_model
        tables = model.tables if hasattr(model, 'tables') else []

        if not tables:
            return

        DocxHelpers.add_heading(self.doc, 'Resumen de Tablas', level=2)

        # Prepare data
        data = []
        for table in tables:
            row = [
                getattr(table, 'name', 'Desconocido'),
                getattr(table, 'table_type', 'Desconocido'),
                str(len(table.columns)) if hasattr(table, 'columns') else '0',
                str(sum(1 for c in table.columns if hasattr(c, 'is_calculated') and c.is_calculated))
                if hasattr(table, 'columns') else '0',
                'Sí' if getattr(table, 'is_hidden', False) else 'No',
                getattr(table, 'description', '')[:50] if getattr(table, 'description', '') else '-'
            ]
            data.append(row)

        headers = ['Nombre de Tabla', 'Tipo', 'Columnas', 'Calculadas', 'Oculta', 'Descripción']
        DocxHelpers.add_table(self.doc, data, headers, style='Table Grid')

    def _add_detailed_tables(self) -> None:
        """Add detailed table documentation"""
        if not hasattr(self.metadata, 'data_model') or not self.metadata.data_model:
            return

        model = self.metadata.data_model
        tables = model.tables if hasattr(model, 'tables') else []

        if not tables:
            return

        DocxHelpers.add_heading(self.doc, 'Documentación Detallada de Tablas', level=2)

        for table in tables:
            self._document_table(table)

    def _document_table(self, table: Any) -> None:
        """
        Document a single table

        Args:
            table: Table object
        """
        # Table name heading
        table_name = getattr(table, 'name', 'Desconocido')
        DocxHelpers.add_heading(self.doc, table_name, level=3)

        # Table properties
        properties = []
        if hasattr(table, 'table_type'):
            properties.append(f"Tipo: {table.table_type}")
        if hasattr(table, 'is_hidden') and table.is_hidden:
            properties.append("Visibilidad: Oculta")
        else:
            properties.append("Visibilidad: Visible")
        if hasattr(table, 'description') and table.description:
            properties.append(f"Descripción: {table.description}")

        if properties:
            DocxHelpers.add_bulleted_list(self.doc, properties)

        # Columns table
        if hasattr(table, 'columns') and table.columns:
            DocxHelpers.add_heading(self.doc, 'Columnas', level=4)

            data = []
            for column in table.columns:
                row = [
                    getattr(column, 'name', 'Desconocido'),
                    getattr(column, 'data_type', 'Desconocido'),
                    'Sí' if getattr(column, 'is_key', False) else 'No',
                    'Sí' if getattr(column, 'is_calculated', False) else 'No',
                    getattr(column, 'description', '')[:80] if getattr(column, 'description', '') else '-'
                ]
                data.append(row)

            headers = ['Nombre de Columna', 'Tipo de Dato', 'Clave', 'Calculada', 'Descripción']
            DocxHelpers.add_table(self.doc, data, headers, style='Table Grid')

        # Hierarchies
        if hasattr(table, 'hierarchies') and table.hierarchies:
            DocxHelpers.add_heading(self.doc, 'Jerarquías', level=4)
            for hierarchy in table.hierarchies:
                hierarchy_name = getattr(hierarchy, 'name', 'Desconocido')
                levels = getattr(hierarchy, 'levels', [])
                level_names = ' → '.join([getattr(level, 'name', 'Desconocido') for level in levels])
                DocxHelpers.add_paragraph(self.doc, f"**{hierarchy_name}**: {level_names}")

        # Source expression for calculated tables
        if hasattr(table, 'source_expression') and table.source_expression:
            DocxHelpers.add_heading(self.doc, 'Expresión de Origen (DAX)', level=4)
            DocxHelpers.add_code_block(self.doc, table.source_expression, language='dax')

        # Add spacing
        self.doc.add_paragraph()
