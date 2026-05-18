"""
Appendix Section Generator

Genera apéndice con código Power Query, parámetros y metadatos.
"""

from docx import Document
from datetime import datetime
from typing import Any
import logging

from ..utils.docx_helpers import DocxHelpers

logger = logging.getLogger(__name__)


class AppendixSectionGenerator:
    """Genera sección de apéndice"""

    def __init__(self, doc: Document, metadata: Any, template_handler: Any):
        """
        Initialize appendix section generator

        Args:
            doc: Document object
            metadata: ReportMetadata object
            template_handler: TemplateHandler object
        """
        self.doc = doc
        self.metadata = metadata
        self.template_handler = template_handler

    def generate(self) -> None:
        """Generate appendix"""
        logger.info("Generating appendix")

        DocxHelpers.add_heading(self.doc, 'Apéndice', level=1)

        # Power Query queries
        self._add_power_query()

        # Report parameters
        self._add_parameters()

        # Extraction metadata
        self._add_extraction_metadata()

        # Glossary
        self._add_glossary()

        logger.info("Appendix generated successfully")

    def _add_power_query(self) -> None:
        """Add Power Query (M) queries"""
        if not hasattr(self.metadata, 'data_model') or not self.metadata.data_model:
            return

        model = self.metadata.data_model

        # Check for queries in tables
        queries = []
        if hasattr(model, 'tables'):
            for table in model.tables:
                if hasattr(table, 'source_expression') and table.source_expression:
                    # Check if it's M code (starts with let/section)
                    source = table.source_expression.strip()
                    if source.lower().startswith(('let', 'section')):
                        queries.append({
                            'name': getattr(table, 'name', 'Desconocido'),
                            'expression': source
                        })

        # Check for explicit queries attribute
        if hasattr(model, 'queries') and model.queries:
            for query in model.queries:
                queries.append({
                    'name': getattr(query, 'name', 'Desconocido'),
                    'expression': getattr(query, 'expression', '')
                })

        if not queries:
            return

        DocxHelpers.add_heading(self.doc, 'Código Power Query (M)', level=2)

        overview_text = (
            f"Las siguientes {len(queries)} consultas de Power Query definen transformaciones de datos "
            "y conexiones a fuentes de datos."
        )
        DocxHelpers.add_paragraph(self.doc, overview_text)

        # Document each query
        for query in queries:
            DocxHelpers.add_heading(self.doc, query['name'], level=3)
            DocxHelpers.add_code_block(self.doc, query['expression'], language='m')
            self.doc.add_paragraph()  # Spacing

    def _add_parameters(self) -> None:
        """Add report parameters"""
        if not hasattr(self.metadata, 'parameters') or not self.metadata.parameters:
            return

        parameters = self.metadata.parameters

        DocxHelpers.add_heading(self.doc, 'Parámetros del Reporte', level=2)

        # Create parameters table
        data = []
        for param in parameters:
            name = getattr(param, 'name', 'Desconocido')
            value = getattr(param, 'current_value', getattr(param, 'default_value', ''))
            param_type = getattr(param, 'type', 'Desconocido')
            description = getattr(param, 'description', '')

            row = [name, str(value), param_type, description[:50] if description else '-']
            data.append(row)

        headers = ['Parámetro', 'Valor', 'Tipo', 'Descripción']
        DocxHelpers.add_table(self.doc, data, headers, style='Table Grid')

    def _add_extraction_metadata(self) -> None:
        """Add extraction metadata"""
        DocxHelpers.add_heading(self.doc, 'Metadatos de Generación del Documento', level=2)

        metadata_dict = {}

        # Extraction date
        metadata_dict['Fecha de Generación'] = datetime.now().strftime('%d/%m/%Y %H:%M:%S')

        # Generator version
        metadata_dict['Versión del Generador'] = 'Generador de Documentación Power BI v3.0'

        # File format
        if hasattr(self.metadata, 'file_format'):
            metadata_dict['Formato de Origen'] = self.metadata.file_format
        elif hasattr(self.metadata, 'file_path'):
            file_path = self.metadata.file_path
            if '.pbix' in str(file_path).lower():
                metadata_dict['Formato de Origen'] = 'PBIX'
            elif '.pbip' in str(file_path).lower():
                metadata_dict['Formato de Origen'] = 'PBIP'

        # File path
        if hasattr(self.metadata, 'file_path'):
            metadata_dict['Archivo de Origen'] = str(self.metadata.file_path)

        # Report name
        metadata_dict['Nombre del Reporte'] = self.metadata.report_name

        # Data model complexity
        if hasattr(self.metadata, 'data_model') and self.metadata.data_model:
            model = self.metadata.data_model
            if hasattr(model, 'tables'):
                metadata_dict['Cantidad de Tablas'] = len(model.tables)
            if hasattr(model, 'measures'):
                metadata_dict['Cantidad de Medidas'] = len(model.measures)
            if hasattr(model, 'relationships'):
                metadata_dict['Cantidad de Relaciones'] = len(model.relationships)

        DocxHelpers.add_statistics_table(self.doc, metadata_dict)

    def _add_glossary(self) -> None:
        """Add Power BI terms glossary"""
        DocxHelpers.add_heading(self.doc, 'Glosario', level=2)

        glossary_terms = {
            'DAX': 'Data Analysis Expressions - Lenguaje de fórmulas para cálculos en Power BI',
            'M (Power Query)': 'Lenguaje de fórmulas de Power Query para transformación de datos',
            'Tabla de Hechos': 'Tabla que contiene datos cuantitativos (medidas) para análisis',
            'Tabla de Dimensiones': 'Tabla que contiene atributos descriptivos para filtrado y agrupación',
            'Cardinalidad': 'Multiplicidad de relación entre tablas (Uno, Muchos)',
            'RLS': 'Row-Level Security - Restringe acceso a datos basado en roles de usuario',
            'OLS': 'Object-Level Security - Controla visibilidad de tablas, columnas o medidas',
            'Dirección de Filtro Cruzado': 'Dirección de propagación de filtro en relaciones',
            'Columna Calculada': 'Columna con valores computados usando DAX al refrescar datos',
            'Medida': 'Cálculo dinámico evaluado en contexto de consulta',
            'Transición de Contexto': 'Conversión de contexto de fila a contexto de filtro en DAX',
            'Función Iteradora': 'Función DAX que evalúa expresión para cada fila (ej. SUMX)',
            'Inteligencia de Tiempo': 'Funciones DAX para cálculos basados en fechas (YTD, MTD, etc.)',
            'Esquema Estrella': 'Diseño de modelo de datos con tabla(s) de hechos rodeada de dimensiones',
            'Esquema Copo de Nieve': 'Esquema estrella con tablas de dimensiones normalizadas',
            'Relación Bidireccional': 'Relación que filtra en ambas direcciones'
        }

        # Create glossary table
        data = [[term, definition] for term, definition in glossary_terms.items()]
        headers = ['Término', 'Definición']
        DocxHelpers.add_table(self.doc, data, headers, style='Table Grid')

        # Add reference
        DocxHelpers.add_paragraph(
            self.doc,
            "\nPara más información, visitar Documentación de Microsoft Power BI: https://docs.microsoft.com/es-es/power-bi"
        )
