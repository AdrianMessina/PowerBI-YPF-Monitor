"""
Visualizations Section Generator

Genera documentación del diseño del reporte y visualizaciones.
"""

from docx import Document
from typing import Any
import logging

from ..utils.docx_helpers import DocxHelpers

logger = logging.getLogger(__name__)


class VisualizationsSectionGenerator:
    """Genera sección de documentación de visualizaciones"""

    def __init__(self, doc: Document, metadata: Any, template_handler: Any):
        """
        Initialize visualizations section generator

        Args:
            doc: Document object
            metadata: ReportMetadata object
            template_handler: TemplateHandler object
        """
        self.doc = doc
        self.metadata = metadata
        self.template_handler = template_handler

    def generate(self) -> None:
        """Generate visualizations section"""
        logger.info("Generating visualizations section")

        DocxHelpers.add_heading(self.doc, 'Visualizaciones del Reporte', level=1)

        if not hasattr(self.metadata, 'report_layout') or not self.metadata.report_layout:
            DocxHelpers.add_paragraph(self.doc, "No hay información de diseño del reporte disponible.")
            return

        layout = self.metadata.report_layout

        # Overview
        self._add_overview(layout)

        # Statistics
        self._add_statistics(layout)

        # Page details
        self._add_page_details(layout)

        # Bookmarks
        if hasattr(layout, 'bookmarks') and layout.bookmarks:
            self._add_bookmarks(layout.bookmarks)

        # Custom visuals
        if hasattr(layout, 'custom_visuals') and layout.custom_visuals:
            self._add_custom_visuals(layout.custom_visuals)

        logger.info("Visualizations section generated successfully")

    def _add_overview(self, layout: Any) -> None:
        """Add visualizations overview"""
        pages = layout.pages if hasattr(layout, 'pages') else []

        overview_text = (
            f"El reporte contiene {len(pages)} página(s) con varias visualizaciones "
            "que presentan información de datos a los usuarios. Cada página está documentada debajo con "
            "detalles sobre los visuales, filtros y configuración."
        )
        DocxHelpers.add_paragraph(self.doc, overview_text)

    def _add_statistics(self, layout: Any) -> None:
        """Add report statistics"""
        DocxHelpers.add_heading(self.doc, 'Estadísticas del Reporte', level=2)

        pages = layout.pages if hasattr(layout, 'pages') else []

        # Count visuals
        total_visuals = 0
        visual_types = {}

        for page in pages:
            if hasattr(page, 'visuals'):
                total_visuals += len(page.visuals)

                for visual in page.visuals:
                    visual_type = getattr(visual, 'type', 'Desconocido')
                    visual_types[visual_type] = visual_types.get(visual_type, 0) + 1

        # Create statistics table
        stats = {
            'Total de Páginas': len(pages),
            'Total de Visuales': total_visuals,
            'Promedio de Visuales por Página': f"{total_visuals / len(pages):.1f}" if pages else '0'
        }

        DocxHelpers.add_statistics_table(self.doc, stats)

        # Visual type breakdown
        if visual_types:
            DocxHelpers.add_heading(self.doc, 'Tipos de Visuales', level=3)

            data = [[vtype, str(count)] for vtype, count in sorted(visual_types.items(), key=lambda x: x[1], reverse=True)]
            headers = ['Tipo de Visual', 'Cantidad']
            DocxHelpers.add_table(self.doc, data, headers, style='Table Grid')

    def _add_page_details(self, layout: Any) -> None:
        """Add detailed page documentation"""
        pages = layout.pages if hasattr(layout, 'pages') else []

        if not pages:
            return

        DocxHelpers.add_heading(self.doc, 'Detalles de Páginas', level=2)

        for page in pages:
            self._document_page(page)

    def _document_page(self, page: Any) -> None:
        """
        Document a single page

        Args:
            page: Page object
        """
        # Page name
        page_name = getattr(page, 'name', 'Desconocido')
        DocxHelpers.add_heading(self.doc, page_name, level=3)

        # Page properties
        properties = []

        if hasattr(page, 'display_name'):
            properties.append(f"Nombre de Visualización: {page.display_name}")

        if hasattr(page, 'is_hidden'):
            visibility = 'Oculta' if page.is_hidden else 'Visible'
            properties.append(f"Visibilidad: {visibility}")

        if hasattr(page, 'width') and hasattr(page, 'height'):
            properties.append(f"Tamaño: {page.width} x {page.height}")

        if properties:
            DocxHelpers.add_bulleted_list(self.doc, properties)

        # Page-level filters
        if hasattr(page, 'filters') and page.filters:
            DocxHelpers.add_paragraph(self.doc, f"Filtros a Nivel de Página: {len(page.filters)} activos")

        # Visuals
        if hasattr(page, 'visuals') and page.visuals:
            DocxHelpers.add_heading(self.doc, f'Visuales ({len(page.visuals)})', level=4)

            # Create visuals table
            data = []
            for visual in page.visuals:
                name = getattr(visual, 'name', 'Desconocido')
                visual_type = getattr(visual, 'type', 'Desconocido')
                title = getattr(visual, 'title', '') if hasattr(visual, 'title') else ''

                # Fields used
                fields = []
                if hasattr(visual, 'fields'):
                    fields = [getattr(f, 'name', str(f)) for f in visual.fields[:3]]  # First 3

                fields_str = ', '.join(fields) if fields else '-'
                if hasattr(visual, 'fields') and len(visual.fields) > 3:
                    fields_str += f" (+{len(visual.fields) - 3} más)"

                # Filters
                filters_count = len(visual.filters) if hasattr(visual, 'filters') else 0
                filters_str = f"{filters_count} filtro(s)" if filters_count > 0 else '-'

                row = [name, visual_type, title or '-', fields_str, filters_str]
                data.append(row)

            headers = ['Nombre', 'Tipo', 'Título', 'Campos', 'Filtros']
            DocxHelpers.add_table(self.doc, data, headers, style='Table Grid')

        # Add spacing
        self.doc.add_paragraph()

    def _add_bookmarks(self, bookmarks: list) -> None:
        """Add bookmarks documentation"""
        DocxHelpers.add_heading(self.doc, 'Marcadores', level=2)

        overview_text = (
            f"El reporte contiene {len(bookmarks)} marcador(es). Los marcadores guardan estados "
            "de vista específicos incluyendo selecciones de página, filtros y valores de segmentación."
        )
        DocxHelpers.add_paragraph(self.doc, overview_text)

        # List bookmarks
        bookmark_names = [getattr(b, 'name', 'Desconocido') for b in bookmarks]
        DocxHelpers.add_bulleted_list(self.doc, bookmark_names)

    def _add_custom_visuals(self, custom_visuals: list) -> None:
        """Add custom visuals documentation"""
        DocxHelpers.add_heading(self.doc, 'Visuales Personalizados', level=2)

        overview_text = (
            f"El reporte usa {len(custom_visuals)} visual(es) personalizado(s). "
            "Los visuales personalizados extienden las capacidades de visualización integradas de Power BI."
        )
        DocxHelpers.add_paragraph(self.doc, overview_text)

        # Create custom visuals table
        data = []
        for visual in custom_visuals:
            name = getattr(visual, 'name', 'Desconocido')
            version = getattr(visual, 'version', 'Desconocido')
            provider = getattr(visual, 'provider', 'Desconocido')

            row = [name, version, provider]
            data.append(row)

        headers = ['Nombre del Visual', 'Versión', 'Proveedor']
        DocxHelpers.add_table(self.doc, data, headers, style='Table Grid')
