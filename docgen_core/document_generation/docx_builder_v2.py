"""
Word Document Builder V2
Fills the official corporate template with real PBIP data
NO INVENTA DATOS - Solo usa información extraída del PBIP
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any
import logging

from .intelligent_mapper import IntelligentTemplateMapper

logger = logging.getLogger(__name__)


class DocxBuilderV2:
    """
    Fills official corporate Word template with real Power BI data.
    Uses the existing template structure and inserts content intelligently.
    """

    def __init__(self, template_path: str):
        """
        Initialize builder with template

        Args:
            template_path: Path to official corporate template
        """
        self.template_path = Path(template_path)
        self.doc: Optional[Document] = None
        self.logger = logging.getLogger(__name__)

    def build(self, metadata: Any, user_inputs: Dict[str, Any],
              er_diagram_path: Optional[str] = None,
              progress_callback=None) -> str:
        """
        Build Word document by filling the template

        Args:
            metadata: Parsed PBIP/TMDL metadata
            user_inputs: User-provided fields (objetivo, alcance, etc.)
            er_diagram_path: Path to generated ER diagram image
            progress_callback: Optional callback(step, message)

        Returns:
            Path to generated document
        """
        self.logger.info("Starting document generation with template filling")

        # Load template
        self._progress(progress_callback, 5, "Cargando template corporativo...")
        self._load_template()

        # Generate context
        self._progress(progress_callback, 10, "Generando contexto desde metadata...")
        mapper = IntelligentTemplateMapper(metadata, user_inputs)
        context = mapper.generate_context()

        try:
            # Fill each section
            self._progress(progress_callback, 15, "Rellenando portada...")
            self._fill_portada(context)

            self._progress(progress_callback, 25, "Rellenando versión del documento...")
            self._fill_version_table(context)

            self._progress(progress_callback, 30, "Rellenando objetivo...")
            self._fill_objective(context)

            self._progress(progress_callback, 35, "Rellenando alcance...")
            self._fill_scope(context)

            self._progress(progress_callback, 40, "Rellenando usuarios...")
            self._fill_users(context)

            self._progress(progress_callback, 50, "Rellenando definiciones (KPIs/Medidas)...")
            self._fill_definitions(context)

            self._progress(progress_callback, 60, "Rellenando orígenes de datos...")
            self._fill_data_sources(context)

            self._progress(progress_callback, 70, "Rellenando filtros...")
            self._fill_filters(context)

            self._progress(progress_callback, 80, "Rellenando modelo ER...")
            self._fill_er_model(context, er_diagram_path)

            self._progress(progress_callback, 90, "Rellenando RLS y seguridad...")
            self._fill_rls(context)

            self._progress(progress_callback, 95, "Rellenando anexo...")
            self._fill_appendix(context)

            # Save document
            self._progress(progress_callback, 98, "Guardando documento...")
            output_path = self._get_output_path(context)
            self.doc.save(str(output_path))

            self._progress(progress_callback, 100, "¡Documento generado exitosamente!")
            self.logger.info(f"Document generated: {output_path}")

            return str(output_path)

        except Exception as e:
            self.logger.error(f"Error generating document: {e}", exc_info=True)
            raise

    def _load_template(self):
        """Load the corporate template"""
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {self.template_path}")

        self.doc = Document(str(self.template_path))
        self.logger.info(f"Template loaded: {self.template_path}")

    # === SECTION FILLERS ===

    def _fill_portada(self, context: Dict[str, Any]):
        """Fill cover page with report name"""
        # Replace [Nombre del tablero] placeholders
        nombre_tablero = context['nombre_tablero']

        for paragraph in self.doc.paragraphs:
            if '[Nombre del tablero]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[Nombre del tablero]', nombre_tablero)
                # Make it bold
                for run in paragraph.runs:
                    if nombre_tablero in run.text:
                        run.bold = True

        self.logger.info(f"Filled portada with: {nombre_tablero}")

    def _fill_version_table(self, context: Dict[str, Any]):
        """Fill version history table"""
        # Find the version table (first table in document)
        if len(self.doc.tables) == 0:
            self.logger.warning("No tables found for version history")
            return

        table = self.doc.tables[0]

        # Fill first data row (row 1, row 0 is header)
        if len(table.rows) >= 2:
            versiones = context['versiones']
            if versiones:
                v = versiones[0]
                cells = table.rows[1].cells

                # Replace placeholders
                cells[0].text = v['version']
                cells[1].text = v['fecha']
                cells[2].text = v['autor']
                cells[3].text = v['observaciones']

        self.logger.info("Filled version table")

    def _fill_objective(self, context: Dict[str, Any]):
        """Fill objective section"""
        objetivo = context['objetivo']

        # Find "Objetivo" heading and insert content after it
        self._insert_content_after_heading('Objetivo', objetivo)

    def _fill_scope(self, context: Dict[str, Any]):
        """Fill scope section"""
        alcance = context['alcance']
        self._insert_content_after_heading('Alcance', alcance)

    def _fill_users(self, context: Dict[str, Any]):
        """Fill users section"""
        usuarios = context['usuarios']
        self._insert_content_after_heading('Usuarios', usuarios)

    def _fill_definitions(self, context: Dict[str, Any]):
        """Fill definitions section with measures"""
        medidas_por_categoria = context['medidas_por_categoria']

        if not medidas_por_categoria:
            self.logger.warning("No measures found to document")
            return

        # Find "Definiciones" heading
        heading_para = self._find_heading_paragraph('Definiciones')

        if not heading_para:
            self.logger.warning("Definiciones heading not found")
            return

        # Insert content after heading
        insert_index = self._get_paragraph_index(heading_para) + 1

        # Remove placeholder text
        self._remove_placeholder_after_heading(heading_para)

        # Add measures by category
        for categoria_data in medidas_por_categoria:
            categoria = categoria_data['categoria']
            medidas = categoria_data['medidas']

            # Add category heading
            cat_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
            cat_para.add_run(f"Categoría: {categoria}").bold = True
            cat_para.style = 'Heading 2'
            insert_index += 1

            # Add each measure
            for medida in medidas:
                # Measure name
                nombre_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
                nombre_para.add_run(f"• {medida['nombre']}").bold = True
                insert_index += 1

                # Formula
                if medida['formula']:
                    formula_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
                    formula_para.add_run("  Fórmula DAX: ").bold = True
                    insert_index += 1

                    code_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
                    code_run = code_para.add_run(f"  {medida['formula']}")
                    code_run.font.name = 'Consolas'
                    code_run.font.size = Pt(9)
                    insert_index += 1

                # Format
                if medida['formato']:
                    fmt_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
                    fmt_para.add_run(f"  Formato: {medida['formato']}")
                    insert_index += 1

                # Description
                if medida['descripcion']:
                    desc_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
                    desc_para.add_run(f"  Descripción: {medida['descripcion']}")
                    insert_index += 1

                # Add spacing
                self.doc.paragraphs[insert_index].insert_paragraph_before()
                insert_index += 1

        self.logger.info(f"Filled {len(medidas_por_categoria)} measure categories")

    def _fill_data_sources(self, context: Dict[str, Any]):
        """Fill data sources section"""
        tablas_origenes = context['tablas_origenes']

        if not tablas_origenes:
            self.logger.warning("No data sources found")
            return

        heading_para = self._find_heading_paragraph('Orígenes')
        if not heading_para:
            return

        insert_index = self._get_paragraph_index(heading_para) + 1
        self._remove_placeholder_after_heading(heading_para)

        for tabla_info in tablas_origenes:
            # Table name
            nombre_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
            nombre_para.add_run(f"Tabla: {tabla_info['tabla']}").bold = True
            nombre_para.style = 'Heading 3'
            insert_index += 1

            # Source type
            tipo_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
            tipo_para.add_run(f"Tipo de fuente: {tabla_info['tipo_fuente']}")
            insert_index += 1

            # Mode
            modo_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
            modo_para.add_run(f"Modo de conexión: {tabla_info['modo']}")
            insert_index += 1

            # Query M (if available)
            if tabla_info['query_m']:
                query_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
                query_para.add_run("Power Query M:").bold = True
                insert_index += 1

                code_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
                code_run = code_para.add_run(tabla_info['query_m'])
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(8)
                insert_index += 1

            # Spacing
            self.doc.paragraphs[insert_index].insert_paragraph_before()
            insert_index += 1

        self.logger.info(f"Filled {len(tablas_origenes)} data sources")

    def _fill_filters(self, context: Dict[str, Any]):
        """Fill filters section"""
        filtros_pq = context['filtros_powerquery']
        filtros_dax = context['filtros_dax']

        heading_para = self._find_heading_paragraph('Filtros')
        if not heading_para:
            return

        insert_index = self._get_paragraph_index(heading_para) + 1
        self._remove_placeholder_after_heading(heading_para)

        # Power Query filters
        if filtros_pq:
            pq_heading = self.doc.paragraphs[insert_index].insert_paragraph_before()
            pq_heading.add_run("Filtros en Power Query").bold = True
            pq_heading.style = 'Heading 2'
            insert_index += 1

            for filtro in filtros_pq:
                f_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
                f_para.add_run(f"• {filtro['tabla']}: {filtro['descripcion']}")
                insert_index += 1

        # DAX filters
        if filtros_dax:
            dax_heading = self.doc.paragraphs[insert_index].insert_paragraph_before()
            dax_heading.add_run("Filtros DAX (Tablas Calculadas)").bold = True
            dax_heading.style = 'Heading 2'
            insert_index += 1

            for filtro in filtros_dax:
                f_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
                f_para.add_run(f"• {filtro['tabla']}: {filtro['descripcion']}")
                insert_index += 1

        if not filtros_pq and not filtros_dax:
            no_filters = self.doc.paragraphs[insert_index].insert_paragraph_before()
            no_filters.add_run("No se detectaron filtros a nivel de modelo.")

        self.logger.info("Filled filters section")

    def _fill_er_model(self, context: Dict[str, Any], er_diagram_path: Optional[str]):
        """Fill ER model section"""
        relaciones = context['relaciones']
        stats = context['estadisticas_modelo']

        heading_para = self._find_heading_paragraph('Modelo ER')
        if not heading_para:
            return

        insert_index = self._get_paragraph_index(heading_para) + 1
        self._remove_placeholder_after_heading(heading_para)

        # Add ER diagram if available
        if er_diagram_path and Path(er_diagram_path).exists():
            diag_heading = self.doc.paragraphs[insert_index].insert_paragraph_before()
            diag_heading.add_run("Diagrama de Relaciones").bold = True
            diag_heading.style = 'Heading 2'
            insert_index += 1

            img_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
            img_run = img_para.add_run()
            img_run.add_picture(er_diagram_path, width=Inches(6.0))
            insert_index += 1

        # Statistics
        stats_heading = self.doc.paragraphs[insert_index].insert_paragraph_before()
        stats_heading.add_run("Estadísticas del Modelo").bold = True
        stats_heading.style = 'Heading 2'
        insert_index += 1

        stats_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
        stats_text = f"""
• Total de tablas: {stats['total_tablas']}
• Total de relaciones: {stats['total_relaciones']}
• Relaciones activas: {stats['relaciones_activas']}
• Relaciones bidireccionales: {stats['relaciones_bidireccionales']}
• Total de medidas: {stats['total_medidas']}
""".strip()
        stats_para.add_run(stats_text)
        insert_index += 1

        # Relationships table
        if relaciones:
            rel_heading = self.doc.paragraphs[insert_index].insert_paragraph_before()
            rel_heading.add_run("Detalle de Relaciones").bold = True
            rel_heading.style = 'Heading 2'
            insert_index += 1

            # Create table
            table = self.doc.add_table(rows=1, cols=5)
            # Try to apply style, fallback to default if not available
            try:
                table.style = 'Light Grid Accent 1'
            except KeyError:
                try:
                    table.style = 'Table Grid'
                except KeyError:
                    pass  # Use default table style

            # Headers
            headers = ['Desde', 'Hacia', 'Cardinalidad', 'Dirección', 'Estado']
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Rows
            for rel in relaciones:
                row_cells = table.add_row().cells
                row_cells[0].text = rel['desde']
                row_cells[1].text = rel['hacia']
                row_cells[2].text = rel['cardinalidad']
                row_cells[3].text = rel['direccion']
                row_cells[4].text = rel['estado']

        self.logger.info(f"Filled ER model with {len(relaciones)} relationships")

    def _fill_rls(self, context: Dict[str, Any]):
        """Fill RLS/Security section"""
        roles = context['roles_rls']
        tiene_rls = context['tiene_rls']

        heading_para = self._find_heading_paragraph('Consideraciones técnicas')
        if not heading_para:
            return

        insert_index = self._get_paragraph_index(heading_para) + 1
        self._remove_placeholder_after_heading(heading_para)

        # Refresh frequency
        freq_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
        freq_para.add_run(f"Frecuencia de actualización: {context['frecuencia_actualizacion']}")
        insert_index += 1

        # Connection mode
        mode_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
        mode_para.add_run(f"Modo de conexión principal: {context['modo_conexion']}")
        insert_index += 1

        # Spacing
        self.doc.paragraphs[insert_index].insert_paragraph_before()
        insert_index += 1

        # RLS section
        rls_heading = self.doc.paragraphs[insert_index].insert_paragraph_before()
        rls_heading.add_run("Seguridad a Nivel de Fila (RLS)").bold = True
        rls_heading.style = 'Heading 2'
        insert_index += 1

        if not tiene_rls:
            no_rls = self.doc.paragraphs[insert_index].insert_paragraph_before()
            no_rls.add_run("No se detectaron roles de seguridad configurados.")
        else:
            for role in roles:
                # Role name
                role_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
                role_para.add_run(f"Rol: {role['nombre']}").bold = True
                insert_index += 1

                # Permissions
                for permiso in role['permisos']:
                    perm_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
                    perm_para.add_run(f"  • Tabla: {permiso['tabla']}")
                    insert_index += 1

                    if permiso['filtro']:
                        filter_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
                        filter_run = filter_para.add_run(f"    Filtro: {permiso['filtro']}")
                        filter_run.font.name = 'Consolas'
                        filter_run.font.size = Pt(9)
                        insert_index += 1

                # Spacing
                self.doc.paragraphs[insert_index].insert_paragraph_before()
                insert_index += 1

        self.logger.info("Filled RLS section")

    def _fill_appendix(self, context: Dict[str, Any]):
        """Fill appendix section"""
        columnas = context['columnas_por_tabla']
        jerarquias = context['jerarquias']

        heading_para = self._find_heading_paragraph('Anexo')
        if not heading_para:
            return

        insert_index = self._get_paragraph_index(heading_para) + 1
        self._remove_placeholder_after_heading(heading_para)

        # Hierarchies
        if jerarquias:
            hier_heading = self.doc.paragraphs[insert_index].insert_paragraph_before()
            hier_heading.add_run("Jerarquías Definidas").bold = True
            hier_heading.style = 'Heading 2'
            insert_index += 1

            for h in jerarquias:
                h_para = self.doc.paragraphs[insert_index].insert_paragraph_before()
                niveles = ' → '.join(h['niveles'])
                h_para.add_run(f"• {h['nombre']} ({h['tabla']}): {niveles}")
                insert_index += 1

        self.logger.info("Filled appendix")

    # === HELPER METHODS ===

    def _insert_content_after_heading(self, heading_text: str, content: str):
        """Insert content after a specific heading"""
        heading_para = self._find_heading_paragraph(heading_text)

        if not heading_para:
            self.logger.warning(f"Heading '{heading_text}' not found")
            return

        # Get index of heading paragraph
        insert_index = self._get_paragraph_index(heading_para) + 1

        # Remove placeholder if exists
        self._remove_placeholder_after_heading(heading_para)

        # Insert new content
        new_para = self.doc.paragraphs[insert_index].insert_paragraph_before(content)
        self.logger.info(f"Inserted content after '{heading_text}'")

    def _find_heading_paragraph(self, heading_text: str):
        """Find paragraph containing heading text"""
        for para in self.doc.paragraphs:
            if heading_text.lower() in para.text.lower():
                style = para.style.name if para.style else ''
                if 'Heading' in style or 'Título' in style or 'Ttulo' in style:
                    return para
        return None

    def _get_paragraph_index(self, paragraph) -> int:
        """Get index of paragraph in document"""
        for i, para in enumerate(self.doc.paragraphs):
            if para == paragraph:
                return i
        return -1

    def _remove_placeholder_after_heading(self, heading_para):
        """Remove placeholder text after heading"""
        heading_idx = self._get_paragraph_index(heading_para)

        if heading_idx < 0 or heading_idx + 1 >= len(self.doc.paragraphs):
            return

        next_para = self.doc.paragraphs[heading_idx + 1]

        # Check if it contains placeholder
        if next_para.text.strip().startswith('[') and next_para.text.strip().endswith(']'):
            # Clear the paragraph
            next_para.text = ''

    def _get_output_path(self, context: Dict[str, Any]) -> Path:
        """Generate output path"""
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_name = context['nombre_tablero'].replace(" ", "_")
        report_name = "".join(c for c in report_name if c.isalnum() or c in ('_', '-'))

        filename = f"PowerBI_Doc_{report_name}_{timestamp}.docx"
        return output_dir / filename

    def _progress(self, callback, step: int, message: str):
        """Call progress callback if provided"""
        if callback:
            callback(step, message)
        self.logger.info(f"[{step}%] {message}")
